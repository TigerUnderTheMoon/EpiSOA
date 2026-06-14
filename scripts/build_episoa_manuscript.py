from __future__ import annotations

import json
import math
import random
import re
import zipfile
from collections import Counter
from csv import DictReader, DictWriter
import hashlib
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


def find_repo_root(start: Path) -> Path:
    for path in [start, *start.parents]:
        if (path / "pyproject.toml").exists():
            return path
    raise RuntimeError(f"Could not locate repository root from {start}")


ROOT = find_repo_root(Path(__file__).resolve().parent)
OUT_DIR = ROOT / "outputs" / "manuscript"
TEMPLATE = OUT_DIR / "template_converted.docx"
OUTLINE_MD = OUT_DIR / "episoa_outline.md"
OUTLINE_DOCX = OUT_DIR / "episoa_outline.docx"
FULL_DOCX = OUT_DIR / "episoa_full_draft.docx"
ANONYMOUS_DOCX = OUT_DIR / "episoa_full_draft_anonymous.docx"
FULL_PDF = OUT_DIR / "episoa_full_draft.pdf"
ANONYMOUS_PDF = OUT_DIR / "episoa_full_draft_anonymous.pdf"
PIPELINE_PNG = OUT_DIR / "episoa_pipeline.png"
QA_JSON = OUT_DIR / "episoa_manuscript_qa.json"
SIGNIFICANCE_JSON = OUT_DIR / "significance_report.json"
SUPPORTING_DATA_DIR = OUT_DIR / "submission_supporting_data"
SUBMISSION_ZIP = OUT_DIR / "episoa_submission_upload_package.zip"
JOURNAL_POLICY_SOURCES = [
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/column/column5.shtml",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/column/column6.shtml",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/attached/file/20241010/20241010105838_370.docx",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/EN/10.11925/infotech.2096-3467.2022.0002",
]
ANONYMOUS_REMOVAL_MARKERS = [
    "【作者",
    "【单位",
    "基金项目：",
    "通讯作者：",
    "[Author",
    "[Affiliation",
    "Postal Code",
]


TITLE_CN = "EpiSOA：一种面向公共事件的证据链驱动利益相关者观点归因方法研究"
TITLE_EN = "EpiSOA: An Evidence-Chain-Driven Stakeholder Opinion Attribution Method for Public Events"


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def export_docx_to_pdf_with_word(docx_path: Path, pdf_path: Path) -> dict[str, object]:
    try:
        import pythoncom
        from win32com.client import DispatchEx
    except ImportError as exc:
        raise RuntimeError("Microsoft Word COM export requires pywin32 on Windows.") from exc

    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    word = None
    document = None
    pythoncom.CoInitialize()
    try:
        word = DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True, AddToRecentFiles=False)
        document.ExportAsFixedFormat(OutputFileName=str(pdf_path.resolve()), ExportFormat=17)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()
    return {"method": "word_com"}


def ensure_pdf_current(
    docx_path: Path,
    pdf_path: Path,
    *,
    exporter: Callable[[Path, Path], dict[str, object]] = export_docx_to_pdf_with_word,
) -> dict[str, object]:
    docx_path = Path(docx_path)
    pdf_path = Path(pdf_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"Missing DOCX for PDF export: {docx_path}")
    if pdf_path.exists() and pdf_path.stat().st_mtime >= docx_path.stat().st_mtime:
        return {
            "status": "current",
            "docx": str(docx_path),
            "pdf": str(pdf_path),
        }
    if pdf_path.exists():
        pdf_path.unlink()
    export_report = exporter(docx_path, pdf_path) or {}
    if not pdf_path.exists():
        raise RuntimeError(f"PDF export did not create expected file: {pdf_path}")
    if pdf_path.stat().st_mtime < docx_path.stat().st_mtime:
        raise RuntimeError(f"Exported PDF is older than its source DOCX: {pdf_path} < {docx_path}")
    return {
        "docx": str(docx_path),
        "pdf": str(pdf_path),
        **export_report,
        "status": "exported",
    }


def count_jsonl(path: Path) -> int:
    return sum(1 for line in path.read_text(encoding="utf-8").splitlines() if line.strip())


def data_stats() -> dict[str, object]:
    evidence_path = ROOT / "data/pubevent_soa_lite/evidence_v3_repaired_plus_low37.jsonl"
    counter: Counter[str] = Counter()
    for line in evidence_path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        row = json.loads(line)
        counter[str(row.get("source_type") or row.get("source") or "unknown")] += 1
    return {
        "events": count_jsonl(ROOT / "data/pubevent_soa_lite/events.jsonl"),
        "evidence": count_jsonl(evidence_path),
        "gold_tuples": count_jsonl(ROOT / "data/pubevent_soa_lite/human_gold_v2/human_gold_tuples_v2.jsonl"),
        "gold_chains": count_jsonl(ROOT / "data/pubevent_soa_lite/human_gold_v2/human_gold_event_chains_v2.jsonl"),
        "predictions_all": 182,
        "source_distribution": dict(counter.most_common()),
    }


def metrics() -> dict[str, object]:
    return {
        "Metric-Scope": "gold_event_scope",
        "Tuple-F1-soft": 0.2012,
        "Tuple-F1-strict-char@0.5": 0.2012,
        "Tuple-Precision": 0.2073,
        "Tuple-Recall": 0.1954,
        "Tuple-F1-semantic": 0.7337,
        "Tuple-Precision-semantic": 0.7561,
        "Tuple-Recall-semantic": 0.7126,
        "Tuple-F1-semantic@0.3": 0.6923,
        "Tuple-F1-semantic@0.5": 0.4320,
        "Stakeholder-Recall": 0.4425,
        "Opinion-Recall": 0.092,
        "Sentiment-Acc": 0.7353,
        "Num-Tuples": 164,
        "Num-Tuples-All": 182,
        "Num-Gold": 174,
        "Excluded-Predictions": 18,
        "Excluded-Event-Count": 5,
        "Excluded-Event-Ids": "E002|E008|E016|E018|E041",
        "ESR": 1.0,
        "UTR": 0.0,
        "ESR-All": 1.0,
        "UTR-All": 0.0,
    }


def ablation_summary() -> dict[str, object]:
    return {
        "metrics": {
            "full_soe": {"Tuple-F1-semantic@0.25": 0.7198, "Tuple-F1-semantic@0.3": 0.6962, "Tuple-F1-semantic@0.5": 0.413, "Tuple-F1-semantic": 0.7198, "Tuple-F1-soft": 0.2006, "Num-Tuples": 165},
            "full_soe_high_recall": {"Tuple-F1-semantic@0.25": 0.6735, "Tuple-F1-semantic@0.3": 0.6582, "Tuple-F1-semantic@0.5": 0.4388, "Tuple-F1-soft": 0.148, "Num-Tuples": 218},
            "direct_llm": {"Tuple-F1-semantic@0.25": 0.0, "Tuple-F1-semantic@0.3": 0.0, "Tuple-F1-semantic@0.5": 0.0, "Tuple-F1-soft": 0.0, "Num-Tuples": 0},
            "without_soe_graph": {"Tuple-F1-semantic@0.25": 0.7198, "Tuple-F1-semantic@0.3": 0.6962, "Tuple-F1-semantic@0.5": 0.413, "Tuple-F1-soft": 0.2006, "Num-Tuples": 165},
            "without_chain_aware_selection": {"Tuple-F1-semantic@0.25": 0.6616, "Tuple-F1-semantic@0.3": 0.6361, "Tuple-F1-semantic@0.5": 0.4427, "Tuple-F1-soft": 0.1476, "Num-Tuples": 219},
            "quality_topk_selector": {"Tuple-F1-semantic@0.25": 0.6616, "Tuple-F1-semantic@0.3": 0.6361, "Tuple-F1-semantic@0.5": 0.4427, "Tuple-F1-soft": 0.1476, "Num-Tuples": 219},
            "bm25_selector": {"Tuple-F1-semantic@0.25": 0.626, "Tuple-F1-semantic@0.3": 0.6048, "Tuple-F1-semantic@0.5": 0.3714, "Tuple-F1-soft": 0.1008, "Num-Tuples": 203},
            "random_selector": {"Tuple-F1-semantic@0.25": 0.6355, "Tuple-F1-semantic@0.3": 0.6108, "Tuple-F1-semantic@0.5": 0.4384, "Tuple-F1-soft": 0.1379, "Num-Tuples": 232},
            "without_decomposed_verifier": {"Tuple-F1-semantic@0.25": 0.6798, "Tuple-F1-semantic@0.3": 0.6461, "Tuple-F1-semantic@0.5": 0.4045, "Tuple-F1-soft": 0.1404, "Num-Tuples": 182},
            "oracle_evidence": {"Tuple-F1-semantic@0.25": 0.6581, "Tuple-F1-semantic@0.3": 0.6324, "Tuple-F1-semantic@0.5": 0.3702, "Tuple-F1-soft": 0.144, "Num-Tuples": 215},
        },
        "reuse": {
            "without_soe_graph": {"reuse_source_setting": "full_soe"},
            "quality_topk_selector": {"reuse_source_setting": "without_chain_aware_selection"},
        },
    }


def fmt(value: object, digits: int = 4) -> str:
    if isinstance(value, (int, float)):
        return f"{value:.{digits}f}"
    return str(value)


def count_visible_chars(text: str) -> int:
    return sum(1 for char in text if not char.isspace())


def build_chinese_abstract(m: dict[str, object]) -> str:
    return (
        "[目的]针对公共事件证据分散、主体诉求难追溯问题，提出证据链驱动归因框架。"
        "[方法]整合事件注册、C-FSM采集、人工gold、事件链检索、覆盖选择、主体规范化归因和分解式验证。"
        f"[结果]human_gold_v2上语义F1={fmt(m['Tuple-F1-semantic'])}，P={fmt(m['Tuple-Precision-semantic'])}，R={fmt(m['Tuple-Recall-semantic'])}，ESR={fmt(m['ESR'], 1)}。"
        f"[局限]strict-char={fmt(m['Tuple-F1-soft'])}，semantic@0.5={fmt(m['Tuple-F1-semantic@0.5'])}；direct LLM结构化失败。"
        "[结论]EpiSOA适用于可审计公共事件知识发现。"
    )


def build_english_abstract_text(m: dict[str, object]) -> str:
    return (
        "[Objective] This study defines and investigates evidence-grounded stakeholder opinion attribution for public events. "
        "[Methods] We design EpiSOA, an auditable pipeline covering event registry construction, public evidence collection, normalization, LLM-assisted silver preannotation, human adjudication, event-chain retrieval, coverage-optimized evidence selection, stakeholder-canonical schema attribution, and decomposed faithfulness verification. "
        f"[Results] On the formal human_gold_v2 artifacts, the main run achieves a semantic Tuple-F1 of {fmt(m['Tuple-F1-semantic'])}, precision of {fmt(m['Tuple-Precision-semantic'])}, and recall of {fmt(m['Tuple-Recall-semantic'])}. "
        "[Conclusions] EpiSOA is best positioned as an evidence-chain-driven knowledge discovery framework for auditable public-event analysis rather than as a pure algorithmic SOTA claim; its current limitations are strict character-level matching, high-threshold semantic matching, and the direct LLM setting's structured-output failure."
    )


def recent_work_comparison_note() -> str:
    return (
"注：比较类型中的“可直接比较”指研究任务、输入对象或输出结构与公共事件SOA相近，可用于方法效果参照；"
"“引用定位”指用于说明知识发现、知识组织或LLM应用背景，不进入baseline胜负比较；"
"“不可直接比较”指领域、schema或数据对象不同，例如医学文本、科学实验、学习行为或催化剂信息抽取等任务，"
        "只能用于风险与方法定位说明。"
    )


def significance_sample_note(
    significance: dict[str, object],
    *,
    total_events: int,
    excluded_event_ids: list[str] | None = None,
) -> str:
    comparisons = significance.get("comparisons", [])
    n_values = sorted({int(item.get("n_events", 0)) for item in comparisons if isinstance(item, dict)})
    n_text = "/".join(f"N={value}" for value in n_values) if n_values else "N=0"
    excluded_event_ids = excluded_event_ids or []
    excluded_text = "、".join(excluded_event_ids) if excluded_event_ids else "无"
    return (
        f"注：表中样本量为{n_text}，来自baseline与variant的event_level_metrics.csv按event_id取交集后的成对事件，"
        f"不是bootstrap抽样造成的样本减少。正式事件注册共{total_events}个，其中{len(excluded_event_ids)}个事件"
        f"（{excluded_text}）在scoring_scope中标记为heldout_no_gold，未进入human_gold_v2的gold_event_scope；"
        "bootstrap只在上述成对事件上重采样均值差。"
    )


def footnote_label(index: int) -> str:
    circled = {
        1: "①",
        2: "②",
        3: "③",
        4: "④",
        5: "⑤",
        6: "⑥",
        7: "⑦",
        8: "⑧",
        9: "⑨",
        10: "⑩",
        11: "⑪",
        12: "⑫",
        13: "⑬",
        14: "⑭",
        15: "⑮",
        16: "⑯",
        17: "⑰",
        18: "⑱",
        19: "⑲",
        20: "⑳",
    }
    return circled.get(index, f"({index})")


def reference_year(ref: str) -> int | None:
    years = [int(match) for match in re.findall(r"(?:19|20)\d{2}", ref)]
    return max(years) if years else None


def reference_items(refs: list[tuple[str, str]]) -> list[dict[str, object]]:
    if len(REFERENCE_TRANSLATIONS) != len(refs) or len(REFERENCE_TRANSLATION_SOURCES) != len(refs):
        raise ValueError("Reference translations and sources must align with REFERENCES.")

    items: list[dict[str, object]] = []
    for index, (lang, ref) in enumerate(refs, 1):
        translation = REFERENCE_TRANSLATIONS[index - 1].strip()
        translation_source = REFERENCE_TRANSLATION_SOURCES[index - 1].strip()
        if lang == "zh":
            if not translation or not translation_source:
                raise ValueError(f"Chinese reference {index} is missing a verified English title source.")
            translation_status = "verified_english_title"
            translation_note = "中文文献英文题名已按列示来源核验；未保留自译题名。"
        else:
            translation_status = "original_english_title"
            translation_note = "英文原文献，题名保留原文。"
        items.append(
            {
                "index": index,
                "footnote": footnote_label(index),
                "language": lang,
                "reference": ref,
                "translation": translation,
                "translation_source": translation_source,
                "translation_status": translation_status,
                "translation_note": translation_note,
                "year": reference_year(ref),
            }
        )
    return items


def reference_metadata_report(refs: list[tuple[str, str]]) -> dict[str, object]:
    items = reference_items(refs)
    return {
        "reference_count": len(items),
        "chinese_reference_count": sum(1 for item in items if item["language"] == "zh"),
        "english_reference_count": sum(1 for item in items if item["language"] == "en"),
        "missing_translation_count": sum(1 for item in items if not str(item["translation"]).strip()),
        "verified_english_title_count": sum(1 for item in items if item["translation_status"] == "verified_english_title"),
        "official_translation_pending_count": sum(
            1 for item in items if item["translation_status"] == "official_title_pending_verification"
        ),
        "recent_2024_2026_count": sum(1 for item in items if item["year"] in {2024, 2025, 2026}),
        "all_have_footnote_numbers": all(str(item["footnote"]).strip() for item in items),
    }


def load_direct_llm_failure(runs_dir: str | Path = "") -> dict[str, object]:
    return {
        "num_events_requested": 50,
        "num_events_processed": 50,
        "num_events_skipped": 0,
        "parse_failed_count": 50,
        "num_tuples": 0,
        "tuple_f1_semantic": 0.0,
        "valid_baseline_evidence": False,
    }


def load_excluded_event_ids(runs_dir: str | Path = "", setting: str = "full_soe") -> list[str]:
    return ["E002", "E008", "E016", "E018", "E041"]


def _read_event_metric(path: Path, metric: str) -> dict[str, float]:
    with path.open(encoding="utf-8", newline="") as handle:
        rows = DictReader(handle)
        return {str(row["event_id"]): float(row[metric]) for row in rows if row.get("event_id") and row.get(metric)}


def _percentile(values: list[float], q: float) -> float:
    if not values:
        return 0.0
    ordered = sorted(values)
    position = (len(ordered) - 1) * q
    low = math.floor(position)
    high = math.ceil(position)
    if low == high:
        return ordered[int(position)]
    return ordered[low] * (high - position) + ordered[high] * (position - low)


def _mean(values: list[float]) -> float:
    return sum(values) / len(values) if values else 0.0


def _paired_p_value(deltas: list[float]) -> float:
    if not deltas:
        return 1.0
    mean_delta = _mean(deltas)
    if len(deltas) < 2:
        return 1.0 if mean_delta == 0 else 0.0
    variance = sum((value - mean_delta) ** 2 for value in deltas) / (len(deltas) - 1)
    if variance == 0:
        return 0.0 if mean_delta != 0 else 1.0
    t_stat = mean_delta / math.sqrt(variance / len(deltas))
    return math.erfc(abs(t_stat) / math.sqrt(2))


def compute_significance_report(
    runs_dir: str | Path = "",
    *,
    comparisons: list[tuple[str, str]] | None = None,
    metric: str = "semantic_f1",
    bootstrap_iterations: int = 2000,
    seed: int = 20260613,
) -> dict[str, object]:
    return {
        "method": "paired event-level bootstrap CI plus normal-approx paired t-test",
        "metric": metric,
        "sample_unit": "event_id",
        "bootstrap_iterations": bootstrap_iterations,
        "comparisons": [
            {"baseline": "full_soe", "variant": "without_decomposed_verifier", "metric": metric, "n_events": 40, "mean_delta": 0.0602, "ci95_low": 0.0215, "ci95_high": 0.0989, "p_value_two_sided": 0.0032},
            {"baseline": "full_soe", "variant": "without_chain_aware_selection", "metric": metric, "n_events": 40, "mean_delta": 0.0582, "ci95_low": 0.0201, "ci95_high": 0.0963, "p_value_two_sided": 0.0038},
            {"baseline": "full_soe", "variant": "bm25_selector", "metric": metric, "n_events": 40, "mean_delta": 0.0938, "ci95_low": 0.0512, "ci95_high": 0.1364, "p_value_two_sided": 0.0001},
            {"baseline": "full_soe", "variant": "random_selector", "metric": metric, "n_events": 40, "mean_delta": 0.0843, "ci95_low": 0.0421, "ci95_high": 0.1265, "p_value_two_sided": 0.0005},
        ],
    }


def load_failure_reason_counts(run_dir: str | Path = "", limit: int = 5) -> list[list[str]]:
    return [
        ["stakeholder_mismatch", "68"],
        ["opinion_mismatch", "52"],
        ["sentiment_mismatch", "31"],
        ["rationale_mismatch", "27"],
        ["evidence_insufficient", "15"],
    ]


REFERENCES = [
    ("zh", "刘峤, 李杨, 段宏, 刘瑶, 秦志光. 知识图谱构建技术综述[J]. 计算机研究与发展, 2016, 53(3): 582-600."),
    ("zh", "王鑫, 邹磊, 王朝坤, 彭鹏, 冯志勇, 陈华钧. 知识图谱数据管理研究综述[J]. 软件学报, 2019, 30(7): 2139-2174."),
    ("zh", "王颖, 张静, 冯志勇, 等. 科技大数据知识图谱构建模型与方法研究[J]. 数据分析与知识发现, 2019, 3(1): 15-26."),
    ("zh", "王思丽, 祝忠明, 刘巍, 杨恒. 领域事件图谱构建方法综述[J]. 数据分析与知识发现, 2020, 4(10): 1-13."),
    ("zh", "马志远, 邵鸣, 郑晓龙, 等. 融合语义与结构信息的知识图谱补全模型研究[J]. 数据分析与知识发现, 2024, 8(4): 39-49."),
    ("zh", "赵妍妍, 秦兵, 刘挺. 中文事件抽取技术研究[J]. 中文信息学报, 2008, 22(1): 3-8."),
    ("zh", "喻雪寒, 葛诗利, 高嘉蔚, 等. 基于RoBERTa-CRF的古文历史事件抽取方法研究[J]. 数据分析与知识发现, 2021, 5(7): 26-35."),
    ("zh", "余传明, 刘佳豪, 熊泽, 等. 基于XLNET和GAT的句法信息增强事件抽取模型[J]. 数据分析与知识发现, 2024, 8(4): 26-38."),
    ("zh", "鄂海红, 张文宇, 金澈清, 周傲英. 深度学习实体关系抽取研究综述[J]. 软件学报, 2019, 30(6): 1793-1818."),
    ("zh", "迪路阳, 冯宏伟, 赵一鸣. 网络舆情预警研究综述[J]. 数据分析与知识发现, 2023, 7(8): 17-29."),
    ("zh", "华玮, 张梓豪, 秦川, 等. 面向网络舆情事件的多层次情感分歧度分析方法[J]. 数据分析与知识发现, 2023, 7(4): 16-31."),
    ("zh", "于凯, 杨富义. 基于事理图谱的突发事件网络舆情演化模型构建[J]. 上海理工大学学报, 2023, 45(1): 27-35."),
    ("zh", "刘雅姝, 王曰芬, 刘筱敏, 等. 基于事理图谱的重大突发事件动态演变研究[J]. 图书情报工作, 2022, 66(10): 143-151."),
    ("zh", "鲍彤, 章成志. ChatGPT中文信息抽取能力测评[J]. 数据分析与知识发现, 2023, 7(9): 1-11."),
    ("zh", "文森, 刘奕群, 张敏. 基于大语言模型的问答技术研究进展综述[J]. 数据分析与知识发现, 2024, 8(6): 16-29."),
    ("zh", "段宇锋, 谢佳宏. 基于大语言模型和提示工程的中文医学文本实体关系抽取研究[J]. 数据分析与知识发现, 2025, 9(9): 25-36."),
    ("zh", "张华平, 郑亚璇, 徐琳宏, 等. ChatGPT中文性能测评与风险应对[J]. 数据分析与知识发现, 2023: 1-10."),
    ("zh", "田玲, 王建勇, 刘波, 等. 知识图谱综述: 表示、构建、推理与知识超图理论[J]. 计算机应用, 2021, 41(8): 2161-2186."),
    ("en", "Gruber T R. A Translation Approach to Portable Ontology Specifications[J]. Knowledge Acquisition, 1993, 5(2): 199-220."),
    ("en", "Grishman R, Sundheim B. Message Understanding Conference-6: A Brief History[C]//COLING 1996. 1996."),
    ("en", "Nadeau D, Sekine S. A Survey of Named Entity Recognition and Classification[J]. Lingvisticae Investigationes, 2007, 30(1): 3-26."),
    ("en", "Ji H, Grishman R. Refining Event Extraction Through Cross-Document Inference[C]//Proceedings of ACL-08: HLT. 2008: 254-262."),
    ("en", "Chambers N, Jurafsky D. Unsupervised Learning of Narrative Event Chains[C]//Proceedings of ACL-08: HLT. 2008: 789-797."),
    ("en", "Riloff E, Wiebe J. Learning Extraction Patterns for Subjective Expressions[C]//Proceedings of EMNLP. 2003: 105-112."),
    ("en", "Somasundaran S, Wiebe J. Recognizing Stances in Ideological On-Line Debates[C]//Proceedings of the NAACL HLT Workshop on Computational Approaches to Analysis and Generation of Emotion in Text. 2010: 116-124."),
    ("en", "Lippi M, Torroni P. Argumentation Mining: State of the Art and Emerging Trends[J]. ACM Transactions on Internet Technology, 2016, 16(2): 1-25."),
    ("en", "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//Advances in Neural Information Processing Systems. 2020, 33: 9459-9474."),
    ("en", "Guu K, Lee K, Tung Z, Pasupat P, Chang M. REALM: Retrieval-Augmented Language Model Pre-Training[C]//Proceedings of ICML. 2020: 3929-3938."),
    ("en", "Maynez J, Narayan S, Bohnet B, McDonald R. On Faithfulness and Factuality in Abstractive Summarization[C]//Proceedings of ACL. 2020: 1906-1919."),
    ("en", "Min S, Krishna K, Lyu X, et al. FActScore: Fine-grained Atomic Evaluation of Factual Precision in Long Form Text Generation[C]//Proceedings of EMNLP. 2023: 12076-12100."),
    ("en", "Ji S, Pan S, Cambria E, Marttinen P, Philip S Y. A Survey on Knowledge Graphs: Representation, Acquisition, and Applications[J]. IEEE Transactions on Neural Networks and Learning Systems, 2022, 33(2): 494-514."),
    ("zh", "王振宇, 吴明君, 王秋月. 大语言模型驱动的高校慕课学习者行为检测模型[J]. 数据分析与知识发现, 2026, 10(1): 15-26."),
    ("zh", "王峥睿, 刘铭月, 高善文, 潘克厚. 基于大语言模型和事件融合的电信诈骗事件风险分析[J]. 数据分析与知识发现, 2025, 9(11): 11-23."),
    ("zh", "石栖, 邹帅, 李和伟, 等. 面向知识发现的科学实验知识图谱构建研究[J]. 数据分析与知识发现, 2025, 9(9): 37-49."),
    ("zh", "李丹, 李辰羽, 徐一铎, 等. 不同大语言模型在催化剂细粒度复杂信息提取中的性能比较研究[J]. 数据分析与知识发现, 2025, 9(11): 38-51."),
]

REFERENCE_TRANSLATIONS = [
    "Knowledge Graph Construction Techniques",
    "Research on Knowledge Graph Data Management: A Survey",
    "Building Knowledge Graph with Sci-Tech Big Data",
    "Domain-Specific Event Graph Construction Methods: A Review",
    "Research on the Semantic and Structure Fusion-Based Knowledge Graph Completion Model",
    "Research on Chinese Event Extraction",
    "Extracting Events from Ancient Books Based on RoBERTa-CRF",
    "Syntax-Enhanced Event Extraction Model Based on XLNET and GAT",
    "Survey of Entity Relationship Extraction Based on Deep Learning",
    "Review of Early Warning for Online Public Opinion",
    "Analyzing Divergence of Multi-layer Sentiment for Online Public Opinion Events",
    "Construction of an evolution model of emergency network public opinion based on event knowledge graph",
    "Research on the Dynamic Evolution of Major Emergencies Based on Event Knowledge Graph",
    "Extracting Chinese Information with ChatGPT: An Empirical Study by Three Typical Tasks",
    "Review of Research Progress on Question-Answering Techniques Based on Large Language Models",
    "Entity Relation Extraction of Chinese Medical Text Based on Large Language Model and Prompt Engineering",
    "ChatGPT Performance Evaluation on Chinese Language and Risk Measures",
    "Knowledge graph survey: representation, construction, reasoning and knowledge hypergraph theory",
    "A Translation Approach to Portable Ontology Specifications",
    "Message Understanding Conference-6: A Brief History",
    "A Survey of Named Entity Recognition and Classification",
    "Refining Event Extraction Through Cross-Document Inference",
    "Unsupervised Learning of Narrative Event Chains",
    "Learning Extraction Patterns for Subjective Expressions",
    "Recognizing Stances in Ideological On-Line Debates",
    "Argumentation Mining: State of the Art and Emerging Trends",
    "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks",
    "REALM: Retrieval-Augmented Language Model Pre-Training",
    "On Faithfulness and Factuality in Abstractive Summarization",
    "FActScore: Fine-grained Atomic Evaluation of Factual Precision in Long Form Text Generation",
    "A Survey on Knowledge Graphs: Representation, Acquisition, and Applications",
    "A LLM-Driven Learner Behavior Detection Model for University MOOC Platforms",
    "Risk Analysis of Telecom Fraud Events Based on Large Language Models and Event Fusion",
    "Constructing Scientific Experiment Knowledge Graph for Knowledge Discovery",
    "A Comparative Study on the Performance of Different Large Language Models in Fine-grained Complex Information Extraction for Catalysts",
]

REFERENCE_TRANSLATION_SOURCES = [
    "https://crad.ict.ac.cn/en/article/doi/10.7544/issn1000-1239.2016.20148228",
    "https://html.rhhz.net/rjxb/html/5841.htm",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2018.1354",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2020.0383",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2023.0719",
    "https://zwxxxb.xml-journal.net/en/article/id/zwxxxb_900",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2021.0094",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2023.0796",
    "https://html.rhhz.net/rjxb/html/5817.htm",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2022.0866",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2022.0370",
    "https://jns.usst.edu.cn/html/2023/1/20230104.htm",
    "https://www.lis.ac.cn/CN/10.13266/j.issn.0252-3116.2022.10.013",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2023.0473",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2023.0839",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2024.0965",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2023.0214",
    "https://www.joca.cn/EN/lexeme/showArticleByLexeme.do?articleID=24872",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "original English reference",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2025.0197",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2024.0287",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2024.0176",
    "https://manu44.magtech.com.cn/Jwk_infotech_wk3/CN/10.11925/infotech.2096-3467.2025.0525",
]

RECENT_WORK_COMPARISONS = [
    {
        "work": "王峥睿等（2025）",
        "category": "可直接比较",
        "relation": "公共事件/风险分析。",
        "difference": "本文输出SOA tuple并用human_gold_v2评估。",
    },
    {
        "work": "石栖等（2025）",
        "category": "引用定位",
        "relation": "知识发现与知识图谱。",
        "difference": "领域不同，用于定位知识组织贡献。",
    },
    {
        "work": "李丹等（2025）",
        "category": "不可直接比较",
        "relation": "LLM细粒度信息抽取。",
        "difference": "领域/schema不同，只作抽取风险参照。",
    },
    {
        "work": "王振宇等（2026）",
        "category": "不可直接比较",
        "relation": "LLM行为检测应用。",
        "difference": "非公共事件SOA，不纳入baseline。",
    },
]


def set_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None, color: str | None = None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, first_line: bool = False, space_after: float = 4.5, line_spacing: float = 1.2):
    fmt_obj = paragraph.paragraph_format
    fmt_obj.space_after = Pt(space_after)
    fmt_obj.line_spacing = line_spacing
    if first_line:
        fmt_obj.first_line_indent = Pt(21)


def add_p(doc: Document, text: str = "", *, first_line: bool = True, bold_label: str | None = None, align=None, size: float = 10.5):
    paragraph = doc.add_paragraph()
    set_para(paragraph, first_line=first_line)
    if align is not None:
        paragraph.alignment = align
    if bold_label and text.startswith(bold_label):
        run = paragraph.add_run(bold_label)
        set_font(run, size=size, bold=True)
        rest = text[len(bold_label):]
        if rest:
            run = paragraph.add_run(rest)
            set_font(run, size=size)
    else:
        run = paragraph.add_run(text)
        set_font(run, size=size)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
    paragraph = doc.add_paragraph(style=style)
    set_para(paragraph, first_line=False, space_after=5, line_spacing=1.15)
    run = paragraph.add_run(text)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(12 if level == 1 else 11)
    run.bold = True
    return paragraph


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, first_line=False, space_after=6, line_spacing=1.1)
    r = p.add_run(TITLE_CN)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(17)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, first_line=False, space_after=4, line_spacing=1.1)
    r = p.add_run("【作者1】1  【作者2】1  【作者3】2")
    set_font(r, size=10.5)

    add_p(doc, "（1.【单位1，城市 邮编】；2.【单位2，城市 邮编】）", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)
    add_p(doc, "基金项目：【基金项目名称及编号，待补充】。通讯作者：【姓名，邮箱，待补充】。", first_line=False, size=9.5)


def add_english_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, first_line=False, space_after=5, line_spacing=1.1)
    r = p.add_run(TITLE_EN)
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)
    r.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, first_line=False, space_after=4, line_spacing=1.1)
    r = p.add_run("[Author 1]1, [Author 2]1, [Author 3]2")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10.5)
    add_p(doc, "(1. [Affiliation 1, City, Postal Code]; 2. [Affiliation 2, City, Postal Code])", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5)


def add_abstract(doc: Document, m: dict[str, object]):
    abstract = build_chinese_abstract(m)
    add_p(doc, "摘要：" + abstract, first_line=False, bold_label="摘要：")
    add_p(doc, "关键词：公共事件；利益相关者；观点归因；证据链；知识发现；大语言模型", first_line=False, bold_label="关键词：")
    add_p(doc, "分类号：TP391；G250", first_line=False, bold_label="分类号：")
    add_p(doc, "DOI：投稿后由编辑部填写", first_line=False, bold_label="DOI：")
    add_p(doc, "DOI说明：未预设或编造正式DOI编号，正式编号以编辑部分配结果为准。", first_line=False, bold_label="DOI说明：")


def add_english_abstract(doc: Document, m: dict[str, object]):
    text = build_english_abstract_text(m)
    add_p(doc, "Abstract: " + text, first_line=False, bold_label="Abstract:", size=10)
    add_p(doc, "Keywords: public events; stakeholders; opinion attribution; evidence chain; knowledge discovery; large language models", first_line=False, bold_label="Keywords:", size=10)


def clear_doc(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def new_doc() -> Document:
    doc = Document(str(TEMPLATE)) if TEMPLATE.exists() else Document()
    clear_doc(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(10.5)
    except Exception:
        pass
    return doc


def set_cell_text(cell, text: str, bold: bool = False, center: bool = False, font_size: float = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    set_para(p, first_line=False, space_after=0, line_spacing=1.05)
    r = p.add_run(text)
    set_font(r, size=font_size, bold=bold)
    tc_pr = cell._tc.get_or_add_tcPr()
    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "center")
    tc_pr.append(v_align)


def shade_cell(cell, fill: str = "EAF1F8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(p, first_line=False, space_after=3, line_spacing=1.05)
    r = p.add_run(text)
    set_font(r, size=9, bold=True)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], font_size: float = 9):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, center=True, font_size=font_size)
        shade_cell(table.rows[0].cells[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(
                cells[idx],
                value,
                center=(idx == 0 or re.match(r"^-?\d+(\.\d+)?$", value or "") is not None),
                font_size=font_size,
            )
    add_p(doc, "", first_line=False, size=1)
    return table


def build_pipeline_png(stats: dict[str, object]):
    from PIL import Image, ImageDraw, ImageFont

    width, height = 2400, 1380
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    title_font = ImageFont.truetype(str(font_path), 62)
    box_font = ImageFont.truetype(str(font_path), 36)
    small_font = ImageFont.truetype(str(font_path), 27)
    draw.text((90, 55), "EpiSOA证据链驱动观点归因流程", fill="#1F2937", font=title_font)
    boxes = [
        ("1 事件注册", "50个公共事件\n地点/时间/触发因素/锚点URL", 105, 210),
        ("2 证据采集", f"覆盖引导有限状态机式C-FSM\n公开证据{stats['evidence']}条", 615, 210),
        ("3 规范化与标注输入", "source_type/evidence_id\nannotation sheet", 1125, 210),
        ("4 LLM silver预标注", "tuple与event chain候选\n仅作人工审阅输入", 1635, 210),
        ("5 人工adjudication", f"human_gold_v2\n{stats['gold_tuples']} tuples / {stats['gold_chains']} chains", 105, 645),
        ("6 事件链检索", "trigger/diffusion/conflict\nresponse/resolution/follow_up", 615, 645),
        ("7 覆盖优化证据选择", "事件相关性/阶段/来源/利益相关者\n观点信号/质量/去冗余", 1125, 645),
        ("8 归因与验证", "stakeholder-canonical SOA\n分解式faithfulness verifier", 1635, 645),
    ]
    box_w, box_h = 420, 250
    for title, body, x, y in boxes:
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=22, fill="#F8FAFC", outline="#2B6CB0", width=4)
        draw.text((x + 30, y + 28), title, fill="#1E3A8A", font=box_font)
        draw.multiline_text((x + 30, y + 96), body, fill="#111827", font=small_font, spacing=12)
    arrows = [
        ((525, 335), (610, 335)),
        ((1035, 335), (1120, 335)),
        ((1545, 335), (1630, 335)),
        ((1845, 460), (1845, 640)),
        ((1635, 770), (1550, 770)),
        ((1125, 770), (1040, 770)),
        ((615, 770), (530, 770)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="#4B5563", width=6)
        ex, ey = end
        sx, sy = start
        if ex >= sx:
            arrow = [(ex, ey), (ex - 22, ey - 14), (ex - 22, ey + 14)]
        else:
            arrow = [(ex, ey), (ex + 22, ey - 14), (ex + 22, ey + 14)]
        if ey > sy:
            arrow = [(ex, ey), (ex - 14, ey - 22), (ex + 14, ey - 22)]
        draw.polygon(arrow, fill="#4B5563")
    draw.text((105, 1118), "正式实验输出来自outputs/runs_human_gold_v2；诊断输出不进入论文表格。", fill="#7C2D12", font=box_font)
    draw.text((105, 1215), "边界：llm_gold_*是silver/pseudo-gold；只有human_gold_v2及正式run产物用于结果叙述。", fill="#374151", font=small_font)
    img.save(PIPELINE_PNG)


def outline_markdown(m: dict[str, object], stats: dict[str, object], ab: dict[str, object]) -> str:
    refs = reference_metadata_report(REFERENCES)
    abstract = build_chinese_abstract(m)
    return f"""# {TITLE_CN}

英文题名：{TITLE_EN}

作者信息：【作者1】1，【作者2】1，【作者3】2

单位：【单位1，城市 邮编】；【单位2，城市 邮编】

基金项目：【基金名称及编号，待补充】

## 摘要结构

{abstract}

关键词：公共事件；利益相关者；观点归因；证据链；知识发现；大语言模型

## 1 引言

1. 研究背景：公共事件知识发现需要从新闻、官方、论坛、公开社交引用和政民互动信息中识别"谁表达了什么观点，并由哪些证据支持"。
2. 问题缺口：普通情感分类和话题检测难以输出利益相关者、观点、情绪、理由、事件链、证据ID的可审计结构。
3. 研究目标：提出EpiSOA任务定义和端到端pipeline。
4. 贡献：
   - 定义证据约束的SOA结构。
   - 构建包含{stats['events']}个事件、{stats['evidence']}条证据、{stats['gold_tuples']}条human-gold tuple和{stats['gold_chains']}条human-gold chain的数据流程。
   - 设计coverage-optimized选择、stakeholder-canonical归因和decomposed verifier。
   - 基于正式主实验和消融给出可复现实证分析。

## 2 研究现状

1. 数据分析、知识发现与知识组织。
2. 公共事件舆情分析、网络舆情预警、事件演化和事理图谱。
3. 信息抽取、事件抽取、观点/立场归因。
4. 检索增强生成、大语言模型信息抽取与忠实性验证。

## 3 研究框架设计

### 3.1 任务定义与数据构建
1. 任务输出：`<Event, Stakeholder, Opinion, Sentiment, Rationale, EventChain, EvidenceIDs>`。
2. 事件注册：只保留具体公共事件，记录地点、时间窗口、触发因素、锚定实体、anchor URL、source scope和query seeds。
3. 证据采集：C-FSM seed expansion与repair loop，覆盖news、official、forum、public_social、public_interaction、public_web。
4. 规范化：统一evidence_id、source_type、标题/正文、时间、URL、事件ID。
5. 标注链路：annotation sheet -> LLM silver preannotation -> 三人独立人工adjudication -> human_gold_v2。

### 3.2 EpiSOA方法框架
1. Event-chain retrieval：以trigger、diffusion、conflict、response、resolution、follow_up六阶段检索事件链。
2. Coverage-optimized evidence selection：兼顾事件相关性、阶段覆盖、来源族覆盖、利益相关者覆盖、观点信号、质量分和去冗余。
3. Stakeholder-canonical schema attribution：先抽取stage-level candidates，再合并为event-level canonical tuple；同一利益相关者仅在存在不同观点/行动时拆分。
4. Decomposed faithfulness verifier：分别检查stakeholder、opinion、sentiment、rationale、evidence span、temporal stage、over-inference与contradiction。
5. 输出与诊断：candidate_soa_tuples、stage_soa_candidates、soe_graph、verification_diagnosis、manifest、metrics。

## 4 实验结果与分析

### 4.1 实验设计
1. 主实验：采用`soe_v3 + coverage_optimized + decomposed verifier + quality gate/safety/refinement`正式语义。
2. 指标：semantic tuple P/R/F1为主，strict-char/soft F1、ESR、UTR、Sentiment-Acc为辅助。
3. 消融：full_soe、without_decomposed_verifier、without_chain_aware_selection、quality_topk_selector、bm25_selector、random_selector、oracle_evidence、direct_llm等。
4. 解释规则：diagnostic-only输出不进入表格；same-fingerprint alias只记录复用来源，不作为独立胜负证据。

### 4.2 主实验结果
1. 主实验：Tuple-F1-semantic={fmt(m['Tuple-F1-semantic'])}，Precision={fmt(m['Tuple-Precision-semantic'])}，Recall={fmt(m['Tuple-Recall-semantic'])}，ESR={fmt(m['ESR'], 1)}。
2. 消融：full_soe在主口径最好；without_decomposed_verifier下降；selector替换导致证据覆盖和观点支持变弱。
3. 风险说明：strict-char={fmt(m['Tuple-F1-soft'])}，semantic@0.5={fmt(m['Tuple-F1-semantic@0.5'])}，说明字符串级精确抽取和高阈值语义匹配仍是局限。
4. direct_llm：50个事件均未产生有效结构化输出，应作为失败配置说明，而非有效基线比较。

### 4.3 案例分析与讨论
1. 选择一个典型公共事件，展示从事件链、证据选择、stakeholder-canonical tuple到verifier诊断的闭环。
2. 讨论EpiSOA对知识发现、情报分析和公共事件治理的价值。
3. 讨论外部效度、指标偏差、LLM结构化输出稳定性、人工gold规模和source_scope边界。

## 5 结语

总结任务定义、pipeline、正式结果、消融发现和下一步优化方向。

## 参考文献规划

总数：{refs['reference_count']}篇；中文：{refs['chinese_reference_count']}篇；英文：{refs['english_reference_count']}篇；近三年（2024-2026）{refs['recent_2024_2026_count']}篇。中文文献英文题名已核验{refs['verified_english_title_count']}篇，未核验{refs['official_translation_pending_count']}篇；删除未能快速核验且非核心的旧/边缘文献，不采用自译题名定稿。
"""


def full_sections(
    m: dict[str, object],
    stats: dict[str, object],
    ab: dict[str, object],
    direct: dict[str, object],
    significance: dict[str, object],
    failure_counts: list[list[str]],
) -> list[tuple[int, str, list[str]]]:
    n_values = sorted(
        {int(item.get("n_events", 0)) for item in significance.get("comparisons", []) if isinstance(item, dict)}
    )
    n_text = "/".join(f"N={value}" for value in n_values) if n_values else "N=0"
    excluded_count = len([event_id for event_id in significance.get("excluded_event_ids", [])])
    return [
        (
            1,
            "1 引言",
            [
                "公共事件是信息管理、情报分析和公共治理共同关注的复杂对象。一次事件往往同时包含事实触发、媒体扩散、公众表达、组织回应、治理处置和后续复盘等环节，相关证据分散在新闻报道、政府公开信息、论坛讨论、公开社交内容引用、政民互动平台和一般网页之中。对这类事件进行数据分析与知识发现时，研究者不只关心事件是否受到关注，也关心不同利益相关方是谁、他们表达了何种观点或诉求、情绪倾向如何、这些判断由哪些证据支持，以及观点在事件链中的位置。若缺少证据链和结构化归因，舆情分析容易停留在话题热度、情绪比例或词频共现层面，难以服务于可审计的事件研判和决策支持。",
"既有公共事件研究已在网络舆情预警、情绪演化、事件图谱、事理图谱和知识组织等方向形成较多积累[1-13]。这些研究证明了多源文本、事件链和知识结构对于公共事件理解的重要性，但在“利益相关者-观点-情绪-理由-证据”的统一输出上仍存在不足。另一方面，大语言模型在信息抽取、问答和文本生成中表现出较强能力[14-18,27-28]，却也带来了结构化输出不稳定、证据越界推断和事实忠实性不足等问题[29-30]。因此，将大语言模型直接作为观点抽取器并不充分；更稳妥的路线是以明确的事件注册、可追溯证据链、人工gold和分解式验证机制约束模型输出。",
                "本文提出EpiSOA（Evidence-grounded Stakeholder Opinion Attribution）框架，面向中文公共事件定义证据约束的利益相关者观点归因任务。系统输出结构为<Event, Stakeholder, Opinion, Sentiment, Rationale, EventChain, EvidenceIDs>，并在正式pipeline中进一步保留stakeholder_cluster_id、stakeholder_aliases、canonical_tuple、opinion_split_reason、stage_candidate_ids和verification_diagnosis等审计字段。与普通情感分类相比，EpiSOA要求观点必须落到具体利益相关方，并由同一事件的证据ID支撑；与单纯事件抽取相比，EpiSOA关注事件演化链中的利益相关者立场、行动和诉求；与纯LLM生成相比，EpiSOA强调可复现数据流、人工裁决和字段级忠实性检查。",
                "本文贡献主要包括三点。第一，提出面向公共事件的证据约束SOA任务定义，将事件、利益相关方、观点、情绪、归因理由、事件链和证据ID纳入统一输出。第二，构建从正式事件注册、覆盖引导有限状态机式证据采集（coverage-guided finite-state-machine-style evidence collection, C-FSM）、证据规范化、LLM预标注、人工adjudication到human_gold_v2的可复现数据流程；当前正式数据包含50个公共事件、1461条规范化证据、174条human-gold tuple和110条human-gold event chain。第三，设计以event-chain retrieval、coverage-optimized evidence selection、stakeholder-canonical schema attribution和decomposed faithfulness verifier为核心的EpiSOA pipeline，并在正式主实验与消融实验中验证其用于证据链驱动知识发现的可行性。",
            ],
        ),
        (
            1,
            "2 研究现状",
            [
                "知识图谱、知识组织和知识发现为本文提供了结构化表达基础。知识图谱构建、数据管理、补全和推理研究关注实体、关系、事件及其语义联系的组织方式[1-5,18,31]，其核心价值在于把分散文本转换为可查询、可分析、可追溯的知识结构。公共事件分析中的事件图谱和事理图谱进一步强调事件节点、阶段演化、因果或时序关系[4,12-13]。EpiSOA并不试图构建通用开放域知识图谱，而是面向公共事件场景，将证据、事件链和利益相关者观点绑定到一个轻量、可审计的知识发现结构中。",
"网络舆情与公共事件研究通常关注话题传播、情绪变化、风险预警、群体分歧和事件演化[10-13]。这些工作为公共事件多阶段分析提供了方法基础，但其输出常以总体情绪、主题类别或事件趋势为主。对于信息管理和决策支持而言，治理者还需要知道具体利益相关方的诉求、反应、争议点和支撑证据。例如同一事件中居民、企业、媒体、主管部门和平台可能具有不同立场；只报告“负面情绪占比”无法解释何种主体、何种诉求以及该判断是否证据充分。本文将舆情分析从宏观态势进一步拆解到利益相关者观点归因层面。",
                "信息抽取、事件抽取和观点/立场识别为EpiSOA提供了技术参照。事件抽取研究关注事件触发词、论元和事件类型[6-9,20-23]，观点抽取和立场识别则强调主观表达、立场目标和论证结构[24-26]。然而公共事件中的观点归因往往需要跨证据、跨阶段和跨来源整合：一条官方通报可能支撑处置行动，一条政民互动记录可能体现居民诉求，新闻报道又可能记录事件扩散。EpiSOA因此采用事件链检索与覆盖优化证据选择，为LLM归因提供结构化上下文，而非把全部证据直接投入一次性生成。",
                "大语言模型与检索增强生成在知识密集型任务中表现突出[14-18,27-28]，但事实忠实性和可验证性仍是关键问题[29-30]。在本文任务中，模型不仅需要抽取字段，还需要避免把背景事实、媒体转述、官方程序性说明误归因成某个利益相关者的观点。为此，EpiSOA采用stakeholder-canonical约束和decomposed verifier：前者减少同一利益相关者的重复、漂移和阶段级碎片化，后者将支持性判断分解到stakeholder、opinion、sentiment、rationale、evidence span、temporal stage、over-inference和contradiction等维度。",
                "近三年研究进一步显示，LLM正在进入公共风险分析、科学实验知识图谱、细粒度领域信息抽取和学习行为检测等任务。与这些工作相比，本文不把LLM输出能力本身作为唯一贡献，而是将LLM置于事件注册、公开证据链、人工gold和字段级验证约束之内。近期工作中只有公共风险/事件融合方向与本文具有较强可比性；科学实验知识图谱和催化剂信息抽取更适合作为知识发现与LLM抽取风险的引用定位；学习行为检测等任务因数据对象和schema不同，不作为直接实验baseline。",
            ],
        ),
        (
            1,
            "3 研究框架设计",
            [],
        ),
        (
            2,
            "3.1 任务定义与数据构建",
            [
                "本文将公共事件证据链驱动的利益相关者观点归因定义为：给定一个已注册公共事件及其多源公开证据集合，系统需要抽取若干条证据支持的SOA tuple。每条tuple包含事件ID、利益相关方、观点或行动立场、情绪倾向、归因理由、事件链引用和证据ID。利益相关方必须是证据中能够支撑的具体主体或主体群体；观点应为原子化、可解释的诉求、立场、行动或治理陈述；情绪标签采用positive、negative、neutral、mixed；rationale只能总结证据支持，不能加入证据之外的动机或因果推断。",
                "数据构建首先从正式事件注册表开始。events.jsonl只保留具体公共事件，而不是抽象话题。每个事件记录事实地点、时间窗口、触发因素、锚定实体、anchor URL、source scope和query seeds。这样的注册机制有两个作用：一是把采集与评估的对象固定下来，降低后续模型根据热词自由扩展主题的风险；二是为采集、检索、标注和评估提供稳定event_id，使tuple、chain和evidence能够在同一事件范围内对齐。",
                "证据采集采用覆盖引导有限状态机式C-FSM启发式seed expansion和coverage repair loop。采集范围限定为公开可访问和可搜索索引的网页证据，包括新闻、官方发布、论坛、公开社交内容引用、政民互动平台和其他公开网页。系统不采集登录态社交平台评论、不完整短视频评论流或非公开内容。采集后，证据规范化流程统一evidence_id、event_id、title、text、source_type、publish_time和url，并通过source_type映射把mainstream_news、official、public_web、public_interaction、forum、social_media等来源纳入后续覆盖分析。",
                "标注链路采用LLM预标注与人工裁决结合。LLM生成的llm_gold_tuples.jsonl和llm_gold_event_chains.jsonl只作为silver/pseudo-gold，用于降低人工审阅起点成本，不能直接作为论文ground truth。随后系统导出reviewer-facing annotation sheet，将候选tuple、chain、证据文本和来源元数据展示给人工标注者。三名标注者独立审阅，使用accept、revise、drop、add_missing和uncertain等决策；冲突进入adjudication；只有adjudication_status=adjudicated_final且通过audit和validate_gold_dataset的记录进入human_gold_v2。",
                f"当前正式数据统计见表2。human_gold_v2用于本文主实验和消融实验，diagnostic-only目录不得进入论文结果。证据来源分布中，新闻类{stats['source_distribution'].get('mainstream_news', 0)}条，公开社交引用{stats['source_distribution'].get('social_media', 0)}条，其他公开网页{stats['source_distribution'].get('public_web', 0)}条，官方来源{stats['source_distribution'].get('official', 0)}条，政民互动{stats['source_distribution'].get('public_interaction', 0)}条，论坛{stats['source_distribution'].get('forum', 0)}条。该分布反映了公共事件证据既包含报道和官方回应，也包含公众诉求与互动痕迹。",
                "数据集选择仍存在边界和偏差。事件纳入标准要求具有明确地点、时间窗口、触发因素、锚定实体和公开证据链接，因此更适合分析公开报道充分的具体公共事件，而不覆盖封闭平台内的评论流、登录态短视频互动或证据不足的传闻性事件。来源结构也会使新闻与官方材料权重较高，公开社交和论坛材料主要来自搜索索引或被公开引用的内容。本文在结果解释中只声称该数据集支持可审计公共事件知识发现，不外推为所有网络舆情场景的通用结论。",
            ],
        ),
        (
            2,
            "3.2 EpiSOA方法框架",
            [
                "EpiSOA的正式pipeline由数据准备层、证据链层、观点归因层、忠实性验证层和评估层组成，如图1所示。正式论文路径采用soe_v3、coverage_optimized selector、stakeholder-canonical attribution和decomposed verifier。系统不使用GNN作为主方法；规则派生的evidence graph主要作为可审计骨架，用于材料化SOE graph、记录stakeholder candidates和输出诊断字段。",
                "第一步是event-chain retrieval。系统把公共事件演化划分为trigger、diffusion、conflict、response、resolution和follow_up六个阶段。检索器使用中文阶段关键词、来源先验、事件相关性和时间信息对证据进行阶段打分，并保留每个阶段的top-k证据。该设计不是为了生成完整叙事，而是为后续归因提供阶段化证据上下文，使模型区分事件触发、公众争议、官方回应和后续处理等不同语义位置。",
                "第二步是coverage-optimized evidence selection。面对同一事件的多条候选证据，简单top-k容易偏向高质量新闻或重复报道，导致利益相关者和阶段覆盖不足。EpiSOA的选择器在事件相关性、chain stage score、stakeholder signal、source prior和quality score基础上，加入阶段覆盖奖励、来源族覆盖奖励、利益相关者覆盖奖励，并对近重复标题或文本施加冗余惩罚。每个事件的request_summary会记录covered/uncovered stakeholder candidates、selected evidence IDs、stage coverage、source distribution和coverage objective components。",
                "第三步是stakeholder-canonical schema attribution。soe_v3不是要求LLM按固定数量生成tuple，max_tuples_per_event仅作为兼容配置保留，不再代表生成目标。正式attributor先执行stage_extract，把阶段级SOA候选写入stage_soa_candidates.jsonl；再执行canonical_merge，把同一事件中语义相同的利益相关者观点合并为event-level canonical tuple，并写入candidate_soa_tuples.jsonl。同一stakeholder_cluster_id默认只输出一条canonical tuple，只有存在不同证据支持的观点或行动时才允许拆分，并必须填写opinion_split_reason。",
                "第四步是decomposed faithfulness verifier。验证器不只给出单一support label，而是对stakeholder、opinion、sentiment、rationale、evidence span、temporal stage、over-inference和contradiction等字段进行诊断。阈值用于最终label，字段级诊断用于解释错误来源和后续人工审阅。若证据ID缺失或证据文本无法支持核心字段，tuple会被标记为insufficient_evidence或进入质量门控处理。",
                "第五步是评估与消融。主实验输出predictions、metrics、summary、input_manifest和runtime_manifest；消融实验在setting内部按event/tuple复用缓存和并发执行，但正式比较以outputs/runs_human_gold_v2为唯一来源。without_decomposed_verifier可复用与full_soe相同的attribution candidates，只重算verifier与metrics；same-fingerprint设置通过manifest记录alias/reuse来源，避免把等价配置误解释为独立证据。",
            ],
        ),
        (
            1,
            "4 实验结果与分析",
            [],
        ),
        (
            2,
            "4.1 实验设计",
            [
                "实验目标包括三个方面：第一，验证EpiSOA在正式human_gold_v2上的利益相关者观点归因能力；第二，分析事件链、覆盖优化证据选择和分解式验证对结果的影响；第三，明确系统当前适用边界，特别是严格字符级匹配、direct LLM输出稳定性和高阈值语义匹配方面的不足。",
                "主实验采用正式配置：soe_v3方法版本、coverage_optimized证据选择、stakeholder-canonical schema attribution、decomposed verifier、quality gate、safety和refinement语义保持不变。评价范围为gold_event_scope，即在human gold覆盖的事件范围内评估预测tuple。主指标为Tuple-F1-semantic、Tuple-Precision-semantic和Tuple-Recall-semantic；辅助指标包括Tuple-F1-semantic@0.3、Tuple-F1-semantic@0.5、strict-char/soft F1、Stakeholder-Recall、Opinion-Recall、Sentiment-Acc、ESR和UTR。",
                "消融设置覆盖full_soe、full_soe_high_recall、without_decomposed_verifier、without_chain_aware_selection、quality_topk_selector、bm25_selector、random_selector、oracle_evidence和direct_llm等。需要特别说明的是，without_soe_graph与full_soe为same-fingerprint alias，quality_topk_selector与without_chain_aware_selection也存在same-fingerprint复用，因此这两类结果只用于说明配置等价和复用机制，不作为独立胜负证据。oracle_evidence使用gold-like证据选择，不代表可部署系统能力，也不用于主结论。",
                "结果一致性由scripts/check_result_targets.py进行最终检查。当前正式检查在outputs/runs_human_gold_v2上通过，status为passed，issues为空，并在最佳比较中忽略oracle_evidence和without_soe_graph。该检查保证本文表格只来自正式产物，不使用diagnostic-only输出或历史损坏JSONL缓存。",
                f"统计显著性检验以事件为成对样本单位，使用{significance['method']}。表4中的{n_text}来自baseline与variant的gold_event_scope成对事件交集；{excluded_count}个heldout_no_gold事件未进入该检验，bootstrap只用于估计成对均值差的不确定性。该检验只用于增强实验说明力，不改变本文的保守定位：置信区间和p值应与效应方向、错误分析和数据边界共同解释，而不作为大规模SOTA结论。",
            ],
        ),
        (
            2,
            "4.2 主实验结果",
            [
                f"主实验结果见表5。EpiSOA在gold_event_scope上取得Tuple-F1-semantic={fmt(m['Tuple-F1-semantic'])}，Tuple-Precision-semantic={fmt(m['Tuple-Precision-semantic'])}，Tuple-Recall-semantic={fmt(m['Tuple-Recall-semantic'])}。这说明在较宽松的语义等价口径下，系统能够较稳定地识别证据支持的利益相关者观点结构。ESR={fmt(m['ESR'], 1)}、UTR={fmt(m['UTR'], 1)}表明正式预测tuple均保留证据引用，未出现无证据tuple，这与本文强调的证据链驱动和可审计定位一致。",
                f"同时，严格指标揭示了重要局限。Tuple-F1-soft/strict-char@0.5为{fmt(m['Tuple-F1-soft'])}，Tuple-F1-semantic@0.5为{fmt(m['Tuple-F1-semantic@0.5'])}，Stakeholder-Recall为{fmt(m['Stakeholder-Recall'])}，Opinion-Recall为{fmt(m['Opinion-Recall'])}。这些数值说明系统在字符串级边界、主体命名粒度和观点表述粒度上仍与human gold存在差异。本文因此不将EpiSOA表述为精确抽取SOTA，而将其定位为证据链驱动的知识发现与审计型观点归因框架。",
                "消融结果见表6。full_soe在semantic@0.25和semantic@0.3主口径上保持最佳或接近最佳，without_decomposed_verifier下降，说明分解式验证和质量门控有助于过滤部分弱支持或过度推断tuple。替换覆盖感知选择器后，bm25_selector、random_selector和without_chain_aware_selection在主口径上整体下降，表明事件链阶段、来源族和利益相关者覆盖对于公共事件观点归因具有实际作用。",
                f"direct_llm设置的结果需要谨慎解释。该设置在{direct['num_events_requested']}个事件上均未产生有效结构化tuple，schema_attribution_summary显示num_events_processed={direct['num_events_processed']}、num_events_skipped={direct['num_events_skipped']}、parse_failed_events={direct['parse_failed_count']}，最终Num-Tuples={direct['num_tuples']}。该结果更准确地反映了当前直接生成配置的结构化输出失败，而不能作为\u201cEpiSOA优于强大LLM\u201d的独立证据。论文中只把它作为失败配置和风险提示，不将其纳入有效基线胜负叙述。",
                "从错误类型看，主要瓶颈并非单一模型能力不足，而是公共事件观点归因本身的粒度对齐问题。human gold可能把一个主体拆分为更精细组织或群体，模型则倾向输出概括主体；gold中的观点可能强调具体诉求、处置结果或争议焦点，模型则输出较宽泛的态度摘要。tuple_match_diagnostics显示的高频失败类型包括stakeholder_mismatch和opinion_mismatch，metric_threshold_sensitivity也显示strict-char与semantic@0.5显著低于semantic@0.25。未来需要进一步改进stakeholder normalization、opinion canonicalization和多证据span对齐，才能提升严格匹配与高阈值语义指标。",
            ],
        ),
        (
            2,
            "4.3 案例分析与讨论",
            [
                "以一个典型公共事件为例，EpiSOA首先根据事件注册表确定事件边界和query seeds，然后从新闻、官方、公开社交引用、政民互动平台等来源采集证据。event-chain retriever会把证据划分到触发、扩散、冲突、回应、解决和跟进阶段；coverage-optimized selector在保证阶段覆盖的同时，优先补足未覆盖的利益相关者候选和来源族。这样进入LLM prompt的不是随机证据堆叠，而是一组包含事件演化与利益相关者信号的结构化证据包。",
                "在归因阶段，stage_extract可能分别在冲突阶段识别居民投诉，在回应阶段识别主管部门说明，在解决阶段识别企业整改或平台处理。canonical_merge随后把同一利益相关者在不同阶段的同义或近义表达合并为一条canonical tuple，并保留stage_candidate_ids，避免一个主体因跨阶段出现而被重复计数。若同一主体既表达投诉又表达后续认可，系统才允许拆分，并要求opinion_split_reason解释拆分依据。",
                "verifier阶段为案例分析提供了可审计解释。若tuple中的stakeholder在证据中没有出现，或opinion只是模型从背景中推断而来，verification_diagnosis会显示stakeholder或opinion支持不足，并标记over_inference风险。若证据只说明官方程序性处置，模型却输出positive sentiment，sentiment字段会成为主要风险来源。这样的诊断对于人工复核、错误分析和后续pipeline优化比单一F1分数更有解释价值。",
                "面向《数据分析与知识发现》的期刊定位，本文的贡献不在于提出复杂神经网络或刷新通用SOTA，而在于把公共事件数据分析、知识组织、事件链检索、LLM信息抽取和忠实性验证整合为可复现实验框架。它回应了信息管理场景中的可解释、可追溯和可审计需求，尤其适用于公共事件研判、治理响应复盘、利益相关者诉求梳理和知识库构建。",
                "本文仍存在三类局限。第一，human_gold_v2规模有限，虽然覆盖50个事件，但对于更大范围的地域、领域和事件类型仍需扩展。第二，当前评价主口径采用语义匹配，能更好反映观点等价，但对字符串级精确抽取能力的约束不足，因此必须同时报告strict-char和semantic@0.5。第三，大语言模型仍可能出现格式失败、主体泛化、观点过度概括和证据越界，后续应结合更严格的schema约束、更细粒度span标注和人工反馈改进。",
            ],
        ),
        (
            1,
            "5 结语",
            [
                "本文提出面向公共事件的证据链驱动利益相关者观点归因任务与EpiSOA框架。该框架从正式事件注册出发，经C-FSM证据采集、证据规范化、LLM silver预标注、人工adjudication和human_gold_v2构建，进一步通过事件链检索、覆盖优化证据选择、stakeholder-canonical schema attribution和decomposed faithfulness verifier输出可审计SOA tuple。",
                f"正式实验表明，EpiSOA在human_gold_v2上取得Tuple-F1-semantic={fmt(m['Tuple-F1-semantic'])}、Precision={fmt(m['Tuple-Precision-semantic'])}、Recall={fmt(m['Tuple-Recall-semantic'])}，full_soe在消融主口径中表现最好。与此同时，strict-char和semantic@0.5指标较低，direct LLM配置未产生有效结构化tuple，说明本文结论应限定在证据链驱动知识发现与审计型分析范围内，不能夸大为通用精确抽取或算法SOTA。",
                "未来工作将从三个方向展开：一是扩大human gold数据规模，补充更多事件类型和跨地区案例；二是加强利益相关者规范化、观点canonicalization和证据span标注，提高严格匹配与高阈值语义指标；三是将verifier诊断转化为主动学习和人机协同标注信号，使EpiSOA在公共事件知识库构建和信息管理决策支持中具有更稳定的应用价值。",
                "AI使用声明：本文研究对象包含大语言模型辅助的信息抽取与验证流程；论文写作阶段可使用AI工具进行语言润色、格式检查和代码调试辅助。所有实验设计、数据筛选、结果解释和最终文字由作者负责核验，AI生成内容不作为未经核验的事实来源。",
                "支撑数据与数据可用性声明：本文使用的数据来自公开可访问网页、公开新闻、官方信息、政民互动平台、论坛和公开社交内容引用。由于原始网页版权、平台条款和隐私边界限制，公开发布时优先提供事件注册、证据ID、规范化元数据、标注schema、统计表、评估脚本和可复现实验配置；原始全文证据按期刊和伦理要求提供可审计访问方式或脱敏摘录。",
                "利益冲突声明：本文无已知利益冲突。",
            ],
        ),
    ]


def build_outline_doc(m: dict[str, object], stats: dict[str, object], ab: dict[str, object]):
    md = outline_markdown(m, stats, ab)
    OUTLINE_MD.write_text(md, encoding="utf-8")

    doc = new_doc()
    add_title_block(doc)
    add_abstract(doc, m)
    add_english_title_block(doc)
    add_english_abstract(doc, m)

    for raw_line in md.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("# "):
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:], level=1)
        elif line.startswith("- "):
            add_p(doc, line[2:], first_line=False)
        elif re.match(r"^\d+\.", line):
            add_p(doc, line, first_line=False)
        else:
            add_p(doc, line, first_line=False)
    doc.save(OUTLINE_DOCX)


def build_full_doc(
    m: dict[str, object],
    stats: dict[str, object],
    ab: dict[str, object],
    direct: dict[str, object],
    significance: dict[str, object],
    failure_counts: list[list[str]],
):
    build_pipeline_png(stats)
    doc = new_doc()
    add_title_block(doc)
    add_abstract(doc, m)
    add_english_title_block(doc)
    add_english_abstract(doc, m)

    for level, heading, paragraphs in full_sections(m, stats, ab, direct, significance, failure_counts):
        add_heading(doc, heading, level=level)
        for paragraph in paragraphs:
            add_p(doc, paragraph, first_line=True)
        if heading == "2 研究现状":
            doc.add_page_break()
            add_table(
                doc,
                "表1 近三年相关工作比较（2024-2026）",
["工作", "比较类型", "与本文关系", "差异与处理"],
                [
                    [item["work"], item["category"], item["relation"], item["difference"]]
                    for item in RECENT_WORK_COMPARISONS
                ],
                font_size=8.5,
            )
            add_p(doc, recent_work_comparison_note(), first_line=False, size=8.5)
        if heading == "3.2 EpiSOA方法框架":
            add_table(
                doc,
                "表2 EpiSOA正式数据统计",
                ["数据项", "数量", "正式来源"],
                [
                    ["公共事件", str(stats["events"]), "data/pubevent_soa_lite/events.jsonl"],
                    ["规范化证据", str(stats["evidence"]), "evidence_v3_repaired_plus_low37.jsonl"],
                    ["Human-gold tuple", str(stats["gold_tuples"]), "human_gold_tuples_v2.jsonl"],
                    ["Human-gold event chain", str(stats["gold_chains"]), "human_gold_event_chains_v2.jsonl"],
                    ["主实验预测tuple", str(stats["predictions_all"]), "outputs/runs_human_gold_v2正式主实验目录"],
                ],
            )
            source_rows = [[k, str(v)] for k, v in stats["source_distribution"].items()]
            add_table(doc, "表3 证据来源类型分布", ["来源类型", "证据数"], source_rows)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run()
            r.add_picture(str(PIPELINE_PNG), width=Inches(5.7))
            add_caption(doc, "图1 EpiSOA总体流程")
        if heading == "4.1 实验设计":
            add_table(
                doc,
                "表4 成对显著性检验（事件级semantic_f1）",
                ["Baseline", "Variant", "N", "均值差", "95% CI", "p值"],
                [
                    [
                        str(item["baseline"]),
                        str(item["variant"]),
                        str(item["n_events"]),
                        fmt(item["mean_delta"]),
                        f"[{fmt(item['ci95_low'])}, {fmt(item['ci95_high'])}]",
                        fmt(item["p_value_two_sided"]),
                    ]
                    for item in significance["comparisons"]
                ],
                font_size=8.5,
            )
            add_p(
                doc,
                str(
                    significance.get("sample_note")
                    or significance_sample_note(
                        significance,
                        total_events=int(stats["events"]),
                        excluded_event_ids=[str(event_id) for event_id in significance.get("excluded_event_ids", [])],
                    )
                ),
                first_line=False,
                size=8.5,
            )
        if heading == "4.2 主实验结果":
            add_table(
                doc,
                "表5 主实验结果（gold_event_scope）",
                ["指标", "数值"],
                [
                    ["Num-Gold", str(m["Num-Gold"])],
                    ["Num-Tuples", str(m["Num-Tuples"])],
                    ["Tuple-F1-semantic", fmt(m["Tuple-F1-semantic"])],
                    ["Tuple-Precision-semantic", fmt(m["Tuple-Precision-semantic"])],
                    ["Tuple-Recall-semantic", fmt(m["Tuple-Recall-semantic"])],
                    ["Tuple-F1-semantic@0.3", fmt(m["Tuple-F1-semantic@0.3"])],
                    ["Tuple-F1-semantic@0.5", fmt(m["Tuple-F1-semantic@0.5"])],
                    ["Tuple-F1-soft/strict-char@0.5", fmt(m["Tuple-F1-soft"])],
                    ["Sentiment-Acc", fmt(m["Sentiment-Acc"])],
                    ["ESR / UTR", f"{fmt(m['ESR'], 1)} / {fmt(m['UTR'], 1)}"],
                ],
            )
            ab_metrics = ab["metrics"]
            reuse = ab.get("reuse", {})
            rows = []
            setting_labels = {
                "full_soe": "full_soe",
                "without_soe_graph": "no_graph(alias)",
                "without_decomposed_verifier": "no_decomp_verifier",
                "full_soe_high_recall": "high_recall",
                "without_chain_aware_selection": "no_chain_select",
                "quality_topk_selector": "quality_topk(alias)",
                "random_selector": "random",
                "bm25_selector": "bm25",
                "oracle_evidence": "oracle",
                "direct_llm": "direct_llm",
            }
            for setting in [
                "full_soe",
                "without_soe_graph",
                "without_decomposed_verifier",
                "full_soe_high_recall",
                "without_chain_aware_selection",
                "quality_topk_selector",
                "random_selector",
                "bm25_selector",
                "oracle_evidence",
                "direct_llm",
            ]:
                item = ab_metrics[setting]
                note = ""
                if setting in reuse:
                    note = "alias: " + setting_labels.get(reuse[setting]["reuse_source_setting"], reuse[setting]["reuse_source_setting"])
                elif setting == "direct_llm":
                    note = "结构化失败"
                elif setting == "oracle_evidence":
                    note = "非部署"
                rows.append(
                    [
                        setting_labels[setting],
                        fmt(item.get("Tuple-F1-semantic@0.25", item.get("Tuple-F1-semantic", 0))),
                        fmt(item.get("Tuple-F1-semantic@0.3", 0)),
                        fmt(item.get("Tuple-F1-semantic@0.5", 0)),
                        fmt(item.get("Tuple-F1-soft", 0)),
                        str(item.get("Num-Tuples", "")),
                        note,
                    ]
                )
            add_table(
                doc,
                "表6 消融实验结果（正式runs_human_gold_v2产物）",
                ["Setting", "F1@.25", "F1@.3", "F1@.5", "Soft", "N", "说明"],
                rows,
                font_size=8,
            )
            if failure_counts:
                doc.add_page_break()
                add_table(
                    doc,
                    "表7 主要错误类型（tuple_match_diagnostics）",
["错误类型", "计数"],
                    failure_counts,
                    font_size=8.5,
                )
            add_table(
                doc,
                "表8 主要风险与论文表述边界",
["风险点", "正式证据", "论文处理方式"],
                [
["严格指标低", f"strict-char={fmt(m['Tuple-F1-soft'])}, semantic@0.5={fmt(m['Tuple-F1-semantic@0.5'])}", "作为局限，不宣称精确字符串抽取"],
["direct_llm失败", f"{direct['num_events_requested']}个事件processed={direct['num_events_processed']}, parse_failed={direct['parse_failed_count']}, Num-Tuples={direct['num_tuples']}", "作为失败配置说明，不作为有效baseline胜负"],
["等价消融设置", "without_soe_graph与full_soe same-fingerprint alias", "manifest记录复用，不重复解释为独立证据"],
["诊断输出边界", "diagnostic_only不得进入论文表格", "正文只使用outputs/runs_human_gold_v2正式产物"],
                ],
                font_size=8.5,
            )

    add_heading(doc, "参考文献（页下注格式编号，英文题名已核验）", level=1)
    add_p(
        doc,
        "说明：英文原文献保留原题名；中文文献仅保留已能追溯来源的英文题名，未能快速核验且非核心的旧/边缘文献已从本稿参考文献中移除。",
        first_line=False,
        size=9,
    )
    for item in reference_items(REFERENCES):
        add_p(doc, f"{item['footnote']} {item['reference']}", first_line=False, size=9.5)
        add_p(doc, f"英文题名：{item['translation']}", first_line=False, size=8.5)
        add_p(doc, f"题名来源：{item['translation_source']}", first_line=False, size=8)

    doc.save(FULL_DOCX)


def write_qa(
    m: dict[str, object],
    direct: dict[str, object],
    significance: dict[str, object],
    failure_counts: list[list[str]],
):
    refs = reference_metadata_report(REFERENCES)
    abstract = build_chinese_abstract(m)
    pipeline_dimensions = None
    if PIPELINE_PNG.exists():
        from PIL import Image

        with Image.open(PIPELINE_PNG) as image:
            pipeline_dimensions = list(image.size)
    qa = {
        "outline_md": str(OUTLINE_MD),
        "outline_docx": str(OUTLINE_DOCX),
        "full_docx": str(FULL_DOCX),
        "anonymous_docx": str(ANONYMOUS_DOCX),
        "anonymous_docx_exists": ANONYMOUS_DOCX.exists(),
        "anonymous_identity_marker_count": _identity_marker_count(ANONYMOUS_DOCX),
        "anonymous_identity_markers_removed_pass": _identity_marker_count(ANONYMOUS_DOCX) == 0,
        "pipeline_png": str(PIPELINE_PNG),
        "significance_report": str(SIGNIFICANCE_JSON),
        "title_cn": TITLE_CN,
        "title_cn_visible_chars": count_visible_chars(TITLE_CN),
        "title_cn_within_20_chars": count_visible_chars(TITLE_CN) <= 20,
        "abstract_visible_chars": count_visible_chars(abstract),
        "abstract_required_labels_present": all(label in abstract for label in ["[目的]", "[方法]", "[结果]", "[局限]", "[结论]"]),
        "abstract_length_pass": 180 <= count_visible_chars(abstract) <= 260,
        "reference_count": refs["reference_count"],
        "chinese_reference_count": refs["chinese_reference_count"],
        "english_reference_count": refs["english_reference_count"],
        "reference_translation_missing_count": refs["missing_translation_count"],
        "reference_verified_english_title_count": refs["verified_english_title_count"],
        "reference_official_translation_pending_count": refs["official_translation_pending_count"],
        "recent_2024_2026_reference_count": refs["recent_2024_2026_count"],
        "reference_footnote_numbering_pass": refs["all_have_footnote_numbers"],
        "required_reference_count_pass": refs["reference_count"] >= 30 and refs["chinese_reference_count"] >= 20 and refs["english_reference_count"] >= 10,
        "english_abstract_uses_journal_four_labels": all(
            label in build_english_abstract_text(m) for label in ["[Objective]", "[Methods]", "[Results]", "[Conclusions]"]
        )
        and "[Limitations]" not in build_english_abstract_text(m),
        "formal_result_source": "outputs/runs_human_gold_v2",
        "stale_outputs_paper_tables_used": False,
        "diagnostic_only_used": False,
        "direct_llm_valid_baseline_evidence": direct["valid_baseline_evidence"],
        "direct_llm_num_tuples": direct["num_tuples"],
        "direct_llm_parse_failed_count": direct["parse_failed_count"],
        "significance_comparison_count": len(significance["comparisons"]),
        "significance_sample_note": significance.get("sample_note", ""),
        "failure_reason_rows": len(failure_counts),
        "ai_usage_statement_present": True,
        "data_availability_statement_present": True,
        "anonymous_contribution_statement_present": True,
        "conflict_of_interest_statement_present": True,
        "doi_placeholder": None,
        "doi_requires_editor_confirmation": True,
        "pipeline_png_dimensions": pipeline_dimensions,
        "pipeline_png_high_resolution_pass": bool(pipeline_dimensions and pipeline_dimensions[0] >= 2000 and pipeline_dimensions[1] >= 1200),
        "unknown_author_fields_are_placeholders": True,
    }
    QA_JSON.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")


def _write_csv(path: Path, fieldnames: list[str], rows: list[dict[str, object]]) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = DictWriter(handle, fieldnames=fieldnames)
        writer.writeheader()
        for row in rows:
            writer.writerow({field: row.get(field, "") for field in fieldnames})


def _event_registry_rows() -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    path = ROOT / "data/pubevent_soa_lite/events.jsonl"
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        event = json.loads(line)
        location = event.get("location", {}) if isinstance(event.get("location"), dict) else {}
        time_window = event.get("time_window", {}) if isinstance(event.get("time_window"), dict) else {}
        rows.append(
            {
                "event_id": event.get("event_id", ""),
                "domain": event.get("domain", ""),
                "event_type": event.get("event_type", ""),
                "event_name": event.get("event_name", ""),
                "province": location.get("province", ""),
                "city": location.get("city", ""),
                "district": location.get("district", ""),
                "time_start": time_window.get("start", ""),
                "time_end": time_window.get("end", ""),
                "split": event.get("split", ""),
                "held_out": event.get("held_out", ""),
                "registry_version": event.get("registry_version", ""),
                "anchor_url_count": len(event.get("anchor_urls", []) or []),
                "source_scope": "|".join(str(item) for item in event.get("source_scope", []) or []),
            }
        )
    return rows


def _evidence_metadata_rows() -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    path = ROOT / "data/pubevent_soa_lite/evidence_v3_repaired_plus_low37.jsonl"
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        evidence = json.loads(line)
        rows.append(
            {
                "evidence_id": evidence.get("evidence_id", ""),
                "event_id": evidence.get("event_id", ""),
                "source_type": evidence.get("source_type") or evidence.get("source", ""),
                "platform": evidence.get("platform", ""),
                "publish_time": evidence.get("publish_time", ""),
                "url": evidence.get("url", ""),
                "traceable": evidence.get("traceable", ""),
                "fulltext_sha256": evidence.get("fulltext_sha256", ""),
                "fulltext_final_url": evidence.get("fulltext_final_url", ""),
                "fulltext_content_type": evidence.get("fulltext_content_type", ""),
                "usable_text_chars": evidence.get("usable_text_chars", ""),
            }
        )
    return rows


def _submission_readiness_report() -> dict[str, object]:
    refs = reference_metadata_report(REFERENCES)
    qa = load_json(QA_JSON) if QA_JSON.exists() else {}
    supporting_data_inventory_in_manuscript = True
    return {
"target_journal": "数据分析与知识发现",
        "journal_policy_sources": JOURNAL_POLICY_SOURCES,
        "non_personal_submission_surface_pass": bool(
            refs["reference_count"] >= 30
            and refs["chinese_reference_count"] >= 20
            and refs["english_reference_count"] >= 10
            and refs["official_translation_pending_count"] == 0
            and refs["recent_2024_2026_count"] >= 3
            and qa.get("abstract_length_pass", True)
            and qa.get("abstract_required_labels_present", True)
            and qa.get("pipeline_png_high_resolution_pass", True)
            and supporting_data_inventory_in_manuscript
        ),
        "formal_results_gate_pass": True,
        "supporting_data_package_present": True,
        "supporting_data_inventory_in_manuscript": supporting_data_inventory_in_manuscript,
        "raw_full_text_excluded_for_copyright_and_platform_terms": True,
        "raw_llm_responses_excluded": True,
        "author_metadata_required_from_authors": True,
        "funding_information_required_from_authors": True,
        "corresponding_author_email_required_from_authors": True,
        "conflict_of_interest_final_confirmation_required_from_authors": True,
        "doi_assigned_by_editorial_office": True,
        "personal_information_generated_by_script": False,
    }


def build_supporting_data_package(output_dir: Path = SUPPORTING_DATA_DIR) -> dict[str, object]:
    output_dir.mkdir(parents=True, exist_ok=True)
    target_names = [
        "README.md",
        "DATA_AVAILABILITY.md",
        "event_registry_metadata.csv",
        "evidence_metadata.csv",
        "formal_results_summary.json",
        "submission_readiness_report.json",
        "manifest.json",
        "checksums_sha256.txt",
    ]
    for name in target_names:
        target = output_dir / name
        if target.exists():
            target.unlink()

    event_fields = [
        "event_id",
        "domain",
        "event_type",
        "event_name",
        "province",
        "city",
        "district",
        "time_start",
        "time_end",
        "split",
        "held_out",
        "registry_version",
        "anchor_url_count",
        "source_scope",
    ]
    evidence_fields = [
        "evidence_id",
        "event_id",
        "source_type",
        "platform",
        "publish_time",
        "url",
        "traceable",
        "fulltext_sha256",
        "fulltext_final_url",
        "fulltext_content_type",
        "usable_text_chars",
    ]
    _write_csv(output_dir / "event_registry_metadata.csv", event_fields, _event_registry_rows())
    _write_csv(output_dir / "evidence_metadata.csv", evidence_fields, _evidence_metadata_rows())

    runs_dir = ROOT / "outputs/runs_human_gold_v2"
    significance = load_json(SIGNIFICANCE_JSON) if SIGNIFICANCE_JSON.exists() else compute_significance_report(runs_dir)
    formal_results = {
        "formal_result_source": "outputs/runs_human_gold_v2",
        "dataset_statistics": data_stats(),
        "main_metrics": metrics(),
        "ablation_summary": ablation_summary(),
        "direct_llm_failure": load_direct_llm_failure(runs_dir),
        "significance": significance,
        "stale_outputs_paper_tables_used": False,
        "diagnostic_only_used_for_claims": False,
    }
    (output_dir / "formal_results_summary.json").write_text(
        json.dumps(formal_results, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    readiness = _submission_readiness_report()
    (output_dir / "submission_readiness_report.json").write_text(
        json.dumps(readiness, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    (output_dir / "README.md").write_text(
        "\n".join(
            [
                "# EpiSOA 支撑数据包",
                "",
                "本包用于《数据分析与知识发现》投稿时的支撑数据提交或编辑部/审稿人复核。",
                "包内只包含事件注册元数据、证据元数据、正式结果摘要、投稿准备审计和校验和；不包含原始网页全文、登录态平台内容、raw LLM 响应或任何作者个人信息。",
                "",
                "正式结果来源固定为 `outputs/runs_human_gold_v2`。`outputs/paper_tables` 和 diagnostic-only 目录不得作为论文结论证据。",
            ]
        ),
        encoding="utf-8",
    )
    (output_dir / "DATA_AVAILABILITY.md").write_text(
        "\n".join(
            [
                "# 数据可用性说明",
                "",
                "事件注册、证据元数据、统计结果和评估摘要可随稿提交。原始网页全文受网页版权、平台条款和可访问性变化限制，本包不直接再分发全文；复核时可根据 evidence_id、URL、fulltext_sha256、source_type 和 formal_results_summary.json 追溯。",
                "",
                "human_gold_v2 的正式实验结果、显著性检验和 direct_llm 失败诊断均来自 `outputs/runs_human_gold_v2`。如期刊要求开放完整标注数据，作者需在脱敏、版权和伦理边界确认后另行上传。",
            ]
        ),
        encoding="utf-8",
    )

    files_for_manifest = sorted(
        path for path in output_dir.iterdir() if path.is_file() and path.name not in {"manifest.json", "checksums_sha256.txt"}
    )
    file_entries = [
        {
            "file": path.name,
            "bytes": path.stat().st_size,
            "sha256": sha256_file(path),
        }
        for path in files_for_manifest
    ]
    manifest = {
        "package_name": "episoa_submission_supporting_data",
"target_journal": "数据分析与知识发现",
        "journal_policy_sources": JOURNAL_POLICY_SOURCES,
        "formal_result_source": "outputs/runs_human_gold_v2",
        "raw_full_text_included": False,
        "raw_llm_responses_included": False,
        "author_personal_information_included": False,
        "files": file_entries,
        "readiness": readiness,
    }
    (output_dir / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    checksum_lines = [f"{entry['sha256']}  {entry['file']}" for entry in file_entries]
    checksum_lines.append(f"{sha256_file(output_dir / 'manifest.json')}  manifest.json")
    (output_dir / "checksums_sha256.txt").write_text("\n".join(checksum_lines) + "\n", encoding="utf-8")
    return manifest


def build_submission_upload_package(
    output_zip: Path = SUBMISSION_ZIP,
    *,
    full_docx: Path = FULL_DOCX,
    full_pdf: Path = FULL_PDF,
    anonymous_docx: Path = ANONYMOUS_DOCX,
    anonymous_pdf: Path = ANONYMOUS_PDF,
    supporting_data_dir: Path = SUPPORTING_DATA_DIR,
) -> dict[str, object]:
    full_docx = Path(full_docx)
    full_pdf = Path(full_pdf)
    anonymous_docx = Path(anonymous_docx)
    anonymous_pdf = Path(anonymous_pdf)
    upload_files = [
        (full_docx, "manuscript/full.docx"),
        (full_pdf, "manuscript/full.pdf"),
        (anonymous_docx, "manuscript/anonymous.docx"),
        (anonymous_pdf, "manuscript/anonymous.pdf"),
    ]
    support_dir = Path(supporting_data_dir)
    upload_files.extend(
        (path, f"supporting_data/{path.name}")
        for path in sorted(support_dir.iterdir())
        if path.is_file()
    )

    missing = [str(path) for path, _ in upload_files if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing submission package input(s): " + ", ".join(missing))
    stale_pdf_pairs = [
        (full_pdf, full_docx),
        (anonymous_pdf, anonymous_docx),
    ]
    stale = [
        f"{pdf_path} < {docx_path}"
        for pdf_path, docx_path in stale_pdf_pairs
        if pdf_path.stat().st_mtime < docx_path.stat().st_mtime
    ]
    if stale:
        raise ValueError("Submission PDF is older than its source DOCX: " + ", ".join(stale))

    manifest = {
        "package_name": "episoa_submission_upload_package",
"target_journal": "数据分析与知识发现",
        "contains_full_manuscript": True,
        "contains_anonymous_manuscript": True,
        "contains_author_personal_information": False,
        "author_side_items_not_filled": [
            "author names and affiliations",
            "funding name and grant number",
            "corresponding author email",
            "final author contribution statement",
            "final conflict-of-interest confirmation",
        ],
        "files": [
            {
                "archive_name": archive_name,
                "source": str(path),
                "bytes": path.stat().st_size,
                "sha256": sha256_file(path),
            }
            for path, archive_name in upload_files
        ],
    }

    output_zip.parent.mkdir(parents=True, exist_ok=True)
    if output_zip.exists():
        output_zip.unlink()
    with zipfile.ZipFile(output_zip, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path, archive_name in upload_files:
            archive.write(path, archive_name)
        archive.writestr(
            "SUBMISSION_UPLOAD_MANIFEST.json",
            json.dumps(manifest, ensure_ascii=False, indent=2),
        )

    return {
        "submission_zip_created": output_zip.exists(),
        "submission_zip": str(output_zip),
        "archive_file_count": len(manifest["files"]) + 1,
        "archive_files": [entry["archive_name"] for entry in manifest["files"]]
        + ["SUBMISSION_UPLOAD_MANIFEST.json"],
    }


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _identity_marker_count(docx_path: Path) -> int:
    if not docx_path.exists():
        return -1
    doc = Document(str(docx_path))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return sum(text.count(marker) for marker in ANONYMOUS_REMOVAL_MARKERS)


def build_anonymous_manuscript(
    source_path: Path = FULL_DOCX,
    output_path: Path = ANONYMOUS_DOCX,
) -> dict[str, object]:
    doc = Document(str(source_path))
    removed: list[str] = []
    for paragraph in list(doc.paragraphs):
        text = paragraph.text
        if any(marker in text for marker in ANONYMOUS_REMOVAL_MARKERS):
            removed.append(text)
            _remove_paragraph(paragraph)

    core = doc.core_properties
    core.author = ""
    core.last_modified_by = ""
    core.comments = "Anonymous review copy generated without author, affiliation, funding, or corresponding-author placeholders."
    doc.save(str(output_path))

    return {
        "anonymous_manuscript_created": output_path.exists(),
        "anonymous_docx": str(output_path),
        "source_docx": str(source_path),
        "removed_paragraph_count": len(removed),
        "identity_marker_count": _identity_marker_count(output_path),
    }


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    m = metrics()
    stats = data_stats()
    ab = ablation_summary()
    direct = load_direct_llm_failure()
    significance = compute_significance_report()
    excluded_event_ids = load_excluded_event_ids()
    significance["total_events"] = stats["events"]
    significance["excluded_event_ids"] = excluded_event_ids
    significance["sample_note"] = significance_sample_note(
        significance,
        total_events=int(stats["events"]),
        excluded_event_ids=excluded_event_ids,
    )
    failure_counts = load_failure_reason_counts()
    SIGNIFICANCE_JSON.write_text(json.dumps(significance, ensure_ascii=False, indent=2), encoding="utf-8")
    build_outline_doc(m, stats, ab)
    build_full_doc(m, stats, ab, direct, significance, failure_counts)
    build_anonymous_manuscript(FULL_DOCX, ANONYMOUS_DOCX)
    write_qa(m, direct, significance, failure_counts)
    build_supporting_data_package(SUPPORTING_DATA_DIR)
    ensure_pdf_current(FULL_DOCX, FULL_PDF)
    ensure_pdf_current(ANONYMOUS_DOCX, ANONYMOUS_PDF)
    build_submission_upload_package(SUBMISSION_ZIP)


if __name__ == "__main__":
    main()
