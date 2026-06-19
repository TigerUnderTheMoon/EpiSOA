from __future__ import annotations

import csv
import importlib
import json
import os
import pytest
import zipfile
from pathlib import Path

from docx import Document


def test_structured_abstract_uses_required_labels_and_compact_length():
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    metrics = {
        "Tuple-F1-semantic": 0.7337,
        "Tuple-Precision-semantic": 0.7561,
        "Tuple-Recall-semantic": 0.7126,
        "Tuple-F1-soft": 0.2012,
        "Tuple-F1-semantic@0.5": 0.4320,
        "ESR": 1.0,
        "UTR": 1.0,
        "Stakeholder-Recall": 0.7,
        "Opinion-Recall": 0.7,
    }

    abstract = builder.build_chinese_abstract(metrics)

    for label in ["[目的]", "[方法]", "[结果]", "[局限]", "[结论]"]:
        assert label in abstract
    assert 180 <= builder.count_visible_chars(abstract) <= 230
    assert "direct LLM" in abstract


def test_english_abstract_uses_journal_four_labels_and_merges_limitations():
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    metrics = {
        "Tuple-F1-semantic": 0.7337,
        "Tuple-Precision-semantic": 0.7561,
        "Tuple-Recall-semantic": 0.7126,
    }

    abstract = builder.build_english_abstract_text(metrics)

    for label in ["[Objective]", "[Methods]", "[Results]", "[Conclusions]"]:
        assert label in abstract
    assert "[Limitations]" not in abstract
    assert "limitations" in abstract.lower()


def test_recent_work_note_explains_non_comparable_category():
    builder = importlib.import_module("scripts.build_episoa_manuscript")

    note = builder.recent_work_comparison_note()

    assert "不可直接比较" in note
    for term in ["领域", "schema", "数据对象"]:
        assert term in note


def test_significance_sample_note_explains_n45_scope_not_bootstrap():
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    significance = {"comparisons": [{"n_events": 45}]}

    note = builder.significance_sample_note(
        significance,
        total_events=50,
        excluded_event_ids=["E002", "E008", "E016", "E018", "E041"],
    )

    assert "N=45" in note
    assert "50" in note
    assert "heldout_no_gold" in note
    assert "不是bootstrap" in note


def test_script_root_points_to_repository_root():
    builder = importlib.import_module("scripts.build_episoa_manuscript")

    assert (builder.ROOT / "pyproject.toml").exists()


def test_direct_llm_failure_summary_uses_schema_artifacts(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    run_dir = tmp_path / "runs"
    direct = run_dir / "ablation_direct_llm"
    direct.mkdir(parents=True)
    (direct / "schema_attribution_summary.json").write_text(
        json.dumps(
            {
                "num_events_requested": 50,
                "num_events_processed": 0,
                "num_events_skipped": 50,
                "parse_failed_events": ["E001", "E002"],
            }
        ),
        encoding="utf-8",
    )
    (direct / "metrics.json").write_text(
        json.dumps({"Num-Tuples": 0, "Tuple-F1-semantic": 0.0}),
        encoding="utf-8",
    )

    summary = builder.load_direct_llm_failure(run_dir)

    assert summary["num_events_requested"] == 50
    assert summary["num_events_processed"] == 0
    assert summary["num_events_skipped"] == 50
    assert summary["parse_failed_count"] == 2
    assert summary["num_tuples"] == 0
    assert summary["valid_baseline_evidence"] is False


def test_direct_llm_valid_baseline_summary_uses_artifacts(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    run_dir = tmp_path / "runs"
    direct = run_dir / "ablation_direct_llm"
    direct.mkdir(parents=True)
    (direct / "schema_attribution_summary.json").write_text(
        json.dumps(
            {
                "num_events_requested": 5,
                "num_events_processed": 5,
                "num_events_skipped": 0,
                "num_api_calls": 6,
                "num_tuples_generated": 9,
                "parse_failed_events": ["E002"],
            }
        ),
        encoding="utf-8",
    )
    (direct / "metrics.json").write_text(
        json.dumps(
            {
                "Metric-Scope": "gold_event_scope",
                "Num-Tuples": 7,
                "Tuple-F1-semantic": 0.42,
            }
        ),
        encoding="utf-8",
    )

    summary = builder.load_direct_llm_failure(run_dir)

    assert summary["num_events_requested"] == 5
    assert summary["num_events_processed"] == 5
    assert summary["num_events_skipped"] == 0
    assert summary["parse_failed_count"] == 1
    assert summary["num_tuples"] == 7
    assert summary["tuple_f1_semantic"] == 0.42
    assert summary["valid_baseline_evidence"] is True


def test_main_metrics_are_loaded_from_formal_artifact(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    runs_dir = tmp_path / "runs"
    main_dir = runs_dir / "pubevent-soa-lite-human-gold-v2-paper"
    main_dir.mkdir(parents=True)
    (main_dir / "metrics.json").write_text(
        json.dumps(
            {
                "Metric-Scope": "gold_event_scope",
                "Tuple-F1-semantic": 0.5123,
                "Tuple-Precision-semantic": 0.6,
                "Tuple-Recall-semantic": 0.45,
                "Tuple-F1-semantic@0.5": 0.5123,
                "Tuple-F1-char@0.5": 0.21,
                "Tuple-F1-soft": 0.21,
                "Tuple-F1-exact": 0.08,
            }
        ),
        encoding="utf-8",
    )

    metrics = builder.metrics(runs_dir)

    assert metrics["Tuple-F1-semantic"] == 0.5123
    assert metrics["Tuple-F1-semantic"] != 0.7337
    assert metrics["Tuple-F1-char@0.5"] == 0.21
    assert metrics["Tuple-F1-exact"] == 0.08


def test_legacy_manuscript_auxiliary_outputs_are_removed(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    for name in builder.LEGACY_AUXILIARY_OUTPUTS:
        (tmp_path / name).write_text("stale 0.7337 strict-char", encoding="utf-8")
    current = tmp_path / "episoa_full_draft.docx"
    current.write_text("current", encoding="utf-8")

    removed = builder.remove_legacy_auxiliary_outputs(tmp_path)

    assert sorted(path.name for path in removed) == sorted(builder.LEGACY_AUXILIARY_OUTPUTS)
    assert current.exists()
    assert not any((tmp_path / name).exists() for name in builder.LEGACY_AUXILIARY_OUTPUTS)


def test_failure_reason_counts_are_loaded_from_failure_audit(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    run_dir = tmp_path / "runs"
    main_dir = run_dir / "pubevent-soa-lite-human-gold-v2-paper"
    main_dir.mkdir(parents=True)
    (main_dir / "tuple_failure_audit.csv").write_text(
        "\n".join(
            [
                "failure_reason,count",
                "opinion_mismatch,7",
                "stakeholder_mismatch,3",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    counts = builder.load_failure_reason_counts(run_dir)

    assert counts == [["opinion_mismatch", "7"], ["stakeholder_mismatch", "3"]]
    assert ["sentiment_mismatch", "31"] not in counts


def test_ablation_summary_updates_direct_llm_from_ablation_results(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    run_dir = tmp_path / "runs"
    run_dir.mkdir()
    (run_dir / "ablation_results.csv").write_text(
        "\n".join(
            [
                "Setting,Metric-Scope,Num-Tuples,Tuple-F1-semantic@0.25,Tuple-F1-semantic@0.3,Tuple-F1-semantic@0.5,Tuple-F1-soft",
                "direct_llm,gold_event_scope,7,0.4100,0.3900,0.2100,0.1200",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    summary = builder.ablation_summary(run_dir)

    assert summary["metrics"]["direct_llm"]["Num-Tuples"] == 7
    assert summary["metrics"]["direct_llm"]["Tuple-F1-semantic@0.25"] == 0.41
    assert summary["metrics"]["direct_llm"]["Tuple-F1-semantic@0.3"] == 0.39
    assert summary["metrics"]["direct_llm"]["Tuple-F1-semantic@0.5"] == 0.21
    assert summary["metrics"]["direct_llm"]["Tuple-F1-soft"] == 0.12


def test_reference_metadata_has_translations_and_recent_work():
    builder = importlib.import_module("scripts.build_episoa_manuscript")

    report = builder.reference_metadata_report(builder.REFERENCES)

    assert report["reference_count"] >= 30
    assert report["missing_translation_count"] == 0
    assert report["recent_2024_2026_count"] >= 3
    assert report["all_have_footnote_numbers"] is True


def test_chinese_references_have_verified_english_titles():
    builder = importlib.import_module("scripts.build_episoa_manuscript")

    items = builder.reference_items(builder.REFERENCES)
    chinese_items = [item for item in items if item["language"] == "zh"]
    english_items = [item for item in items if item["language"] == "en"]

    assert chinese_items
    assert all(item["translation_status"] == "verified_english_title" for item in chinese_items)
    assert all(item["translation_source"].strip() for item in chinese_items)
    assert all(item["translation_status"] == "original_english_title" for item in english_items)


def test_title_block_does_not_render_fake_doi_placeholder():
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    doc = Document()
    metrics = {
        "Tuple-F1-semantic": 0.7337,
        "Tuple-Precision-semantic": 0.7561,
        "Tuple-Recall-semantic": 0.7126,
        "Tuple-F1-soft": 0.2012,
        "Tuple-F1-semantic@0.5": 0.4320,
        "ESR": 1.0,
        "UTR": 1.0,
        "Stakeholder-Recall": 0.7,
        "Opinion-Recall": 0.7,
    }

    builder.add_abstract(doc, metrics)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert "XXXX.XXXX" not in text
    assert "编辑部" in text


def test_supporting_data_package_contains_manifest_and_omits_raw_fulltext(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")

    report = builder.build_supporting_data_package(tmp_path)

    required_files = {
        "README.md",
        "DATA_AVAILABILITY.md",
        "event_registry_metadata.csv",
        "evidence_metadata.csv",
        "formal_results_summary.json",
        "submission_readiness_report.json",
        "checksums_sha256.txt",
        "manifest.json",
    }
    assert required_files <= {path.name for path in tmp_path.iterdir()}
    assert report["raw_full_text_included"] is False
    assert report["formal_result_source"] == "outputs/runs_human_gold_v2"
    assert any("column5.shtml" in source for source in report["journal_policy_sources"])
    assert any("2022.0002" in source for source in report["journal_policy_sources"])
    assert report["readiness"]["non_personal_submission_surface_pass"] is True
    assert report["readiness"]["author_metadata_required_from_authors"] is True
    assert report["readiness"]["supporting_data_inventory_in_manuscript"] is True

    with (tmp_path / "evidence_metadata.csv").open(encoding="utf-8", newline="") as handle:
        reader = csv.DictReader(handle)
        assert "text" not in reader.fieldnames
        assert "legacy_text" not in reader.fieldnames
        assert "usable_text" not in reader.fieldnames
        first_row = next(reader)
    assert first_row["evidence_id"]
    assert first_row["source_type"]


def test_submission_declarations_list_supporting_data_files():
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    metrics = {
        "Tuple-F1-semantic": 0.7337,
        "Tuple-Precision-semantic": 0.7561,
        "Tuple-Recall-semantic": 0.7126,
        "Tuple-F1-soft": 0.2012,
        "Tuple-F1-semantic@0.5": 0.4320,
        "ESR": 1.0,
        "UTR": 1.0,
        "Stakeholder-Recall": 0.7,
        "Opinion-Recall": 0.7,
    }
    stats = {
        "events": 50,
        "evidence": 1461,
        "gold_tuples": 174,
        "gold_chains": 110,
        "predictions_all": 174,
        "source_distribution": {},
    }
    direct = {
        "num_events_requested": 50,
        "num_events_processed": 0,
        "num_events_skipped": 50,
        "num_tuples": 0,
        "parse_failed_count": 50,
        "valid_baseline_evidence": False,
    }
    significance = {
        "method": "paired event-level bootstrap CI plus normal-approx paired t-test",
        "comparisons": [],
        "sample_note": "",
    }

    sections = builder.full_sections(metrics, stats, {}, direct, significance, [])
    declarations = "\n".join(
        "\n".join(paragraphs)
        for _level, heading, paragraphs in sections
        if heading == "9 投稿声明"
    )

    assert "支撑数据清单" in declarations
    for name in ["event_registry_metadata.csv", "evidence_metadata.csv", "formal_results_summary.json"]:
        assert name in declarations


def test_case_study_section_uses_e001_audit_trail():
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    metrics = {
        "Tuple-F1-semantic": 0.7337,
        "Tuple-Precision-semantic": 0.7561,
        "Tuple-Recall-semantic": 0.7126,
        "Tuple-F1-soft": 0.2012,
        "Tuple-F1-semantic@0.5": 0.4320,
        "ESR": 1.0,
        "UTR": 1.0,
        "Stakeholder-Recall": 0.7,
        "Opinion-Recall": 0.7,
    }
    stats = {
        "events": 50,
        "evidence": 1461,
        "gold_tuples": 174,
        "gold_chains": 110,
        "predictions_all": 174,
        "source_distribution": {},
    }
    direct = {
        "num_events_requested": 50,
        "num_events_processed": 0,
        "num_events_skipped": 50,
        "num_tuples": 0,
        "parse_failed_count": 50,
        "valid_baseline_evidence": False,
    }
    significance = {
        "method": "paired event-level bootstrap CI plus normal-approx paired t-test",
        "comparisons": [],
        "sample_note": "",
    }

    sections = builder.full_sections(metrics, stats, {}, direct, significance, [])
    case_text = "\n".join(
        "\n".join(paragraphs)
        for _level, heading, paragraphs in sections
        if heading == "4.3 案例分析与讨论"
    )

    for expected in [
        "E001",
        "ev-00001",
        "ev-00013",
        "stakeholder-canonical",
        "canonical tuple",
        "verification_diagnosis",
        "support_score=1.0",
    ]:
        assert expected in case_text


def test_anonymous_manuscript_removes_author_identifying_placeholders(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    source = tmp_path / "source.docx"
    target = tmp_path / "anonymous.docx"
    doc = Document()
    for text in [
        "证据链驱动观点归因",
        "【作者1】1  【作者2】1  【作者3】2",
        "（1.【单位1，城市 邮编】；2.【单位2，城市 邮编】）",
        "基金项目：【基金项目名称及编号，待补充】。通讯作者：【姓名，邮箱，待补充】。",
        "摘要：[目的]测试摘要。",
        "Evidence-Chain-Driven Opinion Attribution",
        "[Author 1]1, [Author 2]1, [Author 3]2",
        "(1. [Affiliation 1, City, Postal Code]; 2. [Affiliation 2, City, Postal Code])",
        "AI使用声明：作者负责最终核验。",
    ]:
        doc.add_paragraph(text)
    doc.save(source)

    report = builder.build_anonymous_manuscript(source, target)
    text = "\n".join(paragraph.text for paragraph in Document(target).paragraphs)

    assert report["anonymous_manuscript_created"] is True
    assert report["removed_paragraph_count"] >= 4
    for marker in ["【作者", "【单位", "基金项目", "通讯作者", "[Author", "[Affiliation", "Postal Code", "待补充"]:
        assert marker not in text
    assert "证据链驱动观点归因" in text
    assert "摘要：" in text
    assert "AI使用声明" in text


def test_submission_upload_package_contains_only_expected_submission_artifacts(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    full_docx = tmp_path / "full.docx"
    full_pdf = tmp_path / "full.pdf"
    anonymous_docx = tmp_path / "anonymous.docx"
    anonymous_pdf = tmp_path / "anonymous.pdf"
    supporting = tmp_path / "supporting"
    supporting.mkdir()
    for path in [full_docx, full_pdf, anonymous_docx, anonymous_pdf]:
        path.write_text(path.name, encoding="utf-8")
    (supporting / "manifest.json").write_text("{}", encoding="utf-8")
    (supporting / "README.md").write_text("readme", encoding="utf-8")
    package = tmp_path / "submission.zip"

    report = builder.build_submission_upload_package(
        package,
        full_docx=full_docx,
        full_pdf=full_pdf,
        anonymous_docx=anonymous_docx,
        anonymous_pdf=anonymous_pdf,
        supporting_data_dir=supporting,
    )

    assert report["submission_zip_created"] is True
    with zipfile.ZipFile(package) as archive:
        names = set(archive.namelist())
    assert "manuscript/full.docx" in names
    assert "manuscript/full.pdf" in names
    assert "manuscript/anonymous.docx" in names
    assert "manuscript/anonymous.pdf" in names
    assert "supporting_data/manifest.json" in names
    assert "supporting_data/README.md" in names
    assert "SUBMISSION_UPLOAD_MANIFEST.json" in names
    assert all("pdf_pages" not in name and "__pycache__" not in name for name in names)


def test_submission_upload_package_rejects_pdf_older_than_docx(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    full_docx = tmp_path / "full.docx"
    full_pdf = tmp_path / "full.pdf"
    anonymous_docx = tmp_path / "anonymous.docx"
    anonymous_pdf = tmp_path / "anonymous.pdf"
    supporting = tmp_path / "supporting"
    supporting.mkdir()
    for path in [full_docx, full_pdf, anonymous_docx, anonymous_pdf]:
        path.write_text(path.name, encoding="utf-8")
    (supporting / "manifest.json").write_text("{}", encoding="utf-8")
    (supporting / "README.md").write_text("readme", encoding="utf-8")
    os.utime(full_pdf, (1_700_000_000, 1_700_000_000))
    os.utime(full_docx, (1_700_000_100, 1_700_000_100))

    with pytest.raises(ValueError, match="older than its source DOCX"):
        builder.build_submission_upload_package(
            tmp_path / "submission.zip",
            full_docx=full_docx,
            full_pdf=full_pdf,
            anonymous_docx=anonymous_docx,
            anonymous_pdf=anonymous_pdf,
            supporting_data_dir=supporting,
        )


def test_ensure_pdf_current_refreshes_stale_pdf_with_injected_exporter(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    docx = tmp_path / "manuscript.docx"
    pdf = tmp_path / "manuscript.pdf"
    docx.write_text("docx", encoding="utf-8")
    pdf.write_text("old pdf", encoding="utf-8")
    os.utime(pdf, (1_700_000_000, 1_700_000_000))
    os.utime(docx, (1_700_000_100, 1_700_000_100))
    calls = []

    def fake_exporter(source, target):
        calls.append((source, target))
        target.write_text("new pdf", encoding="utf-8")
        os.utime(target, (1_700_000_200, 1_700_000_200))
        return {"method": "fake"}

    report = builder.ensure_pdf_current(docx, pdf, exporter=fake_exporter)

    assert calls == [(docx, pdf)]
    assert report["status"] == "exported"
    assert report["method"] == "fake"
    assert pdf.read_text(encoding="utf-8") == "new pdf"


def test_paired_significance_uses_event_level_metrics(tmp_path):
    builder = importlib.import_module("scripts.build_episoa_manuscript")
    runs = tmp_path / "runs"
    full = runs / "ablation_full_soe"
    weak = runs / "ablation_without_decomposed_verifier"
    full.mkdir(parents=True)
    weak.mkdir(parents=True)

    _write_event_metrics(full / "event_level_metrics.csv", [0.8, 0.7, 0.6])
    _write_event_metrics(weak / "event_level_metrics.csv", [0.4, 0.3, 0.2])

    report = builder.compute_significance_report(
        runs,
        comparisons=[("full_soe", "without_decomposed_verifier")],
        metric="semantic_f1",
        bootstrap_iterations=200,
        seed=7,
    )

    item = report["comparisons"][0]
    assert item["baseline"] == "full_soe"
    assert item["variant"] == "without_decomposed_verifier"
    assert item["n_events"] == 3
    assert item["mean_delta"] > 0
    assert item["ci95_low"] > 0
    assert item["p_value_two_sided"] <= 0.25


def _write_event_metrics(path: Path, values: list[float]) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["event_id", "semantic_f1"])
        writer.writeheader()
        for idx, value in enumerate(values, 1):
            writer.writerow({"event_id": f"E{idx:03d}", "semantic_f1": value})
